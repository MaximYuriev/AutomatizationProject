from fastapi import APIRouter, Body, Depends, HTTPException,UploadFile,File
from fastapi.responses import JSONResponse, FileResponse
from starlette import status
from typing import List, Union, Annotated
from sqlalchemy.orm import Session
from database.db import engine_s
from models.users import *
from models.cases import *
from models.methods import *
import hashlib
import shutil
import os
from api.router import get_session
import docx
from docx.shared import Pt, RGBColor
import pymorphy3
import json
alg_router = APIRouter(prefix='/algorithm')
@alg_router.post("/api/procedural_position",tags=['Procedural Position'])
def create_procedural_position(item:createProceduralPosition,DB: Session = Depends(get_session)):
    procedural_pos = DB.query(ProceduralPosition).filter(ProceduralPosition.name_position == item.name).first()
    if procedural_pos is not None:
        return JSONResponse(status_code= 400, content={"message": "Процессуальное положение уже добавлено"})
    procedural_pos=ProceduralPosition(name_position=item.name)
    try:
        DB.add(procedural_pos)
        DB.commit()
        DB.refresh(procedural_pos)
        return JSONResponse(status_code=200, content={"message": "Процессуальное положение успешно добавлено!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/api/procedural_position",tags=['Procedural Position'])
def get_procedural_position(DB: Session = Depends(get_session)):
    procedural_pos = DB.query(ProceduralPosition).all()
    if procedural_pos is None:
        return JSONResponse(status_code= 404, content={"message": "Процессуальные положения не найдены"})
    return procedural_pos
@alg_router.get("/individual_involved_info/{id}",tags=['Person Involved'])
def get_individual_involved_info(id:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == id).first()
    if cases is None:
        return JSONResponse(status_code=404, content={"message":"Дело не найдено!"})
    individual_id = DB.query(Statement).filter(Statement.statement_id == cases.FK_statement_id).first().fk_individuals_id
    individual = DB.query(Invidivual).filter(Invidivual.individuals_id==individual_id).first()
    passport = DB.query(Passport).filter(Passport.passports_id==individual.FK_passports_id).first()
    info = {
        "first_name":individual.first_name,
        "middle_name":individual.middle_name,
        "last_name":individual.last_name,
        "residence":passport.place_of_residence,
        "passport_serial":passport.passport_serial,
        "passport_number":passport.passport_number
    }
    return info
@alg_router.post("/person_involved/",tags=['Person Involved'])
def create_person_involved(item:createPersonInvolved,DB:Session = Depends(get_session)):
    passport = DB.query(Passport).filter(item.passport_serial == Passport.passport_serial,item.passport_number==Passport.passport_number).first()
    if passport is not None:
        person = DB.query(Invidivual).filter(Invidivual.FK_passports_id == passport.passports_id,Invidivual.first_name==item.first_name,Invidivual.middle_name==item.middle_name,Invidivual.last_name==item.last_name).first()
        if person is not None:
            return person.individuals_id
        else:
            return JSONResponse(status_code=400,content={"message":"Данный паспорт зарегистрирован в системе!"})
    passport = Passport(passport_serial=item.passport_serial,passport_number = item.passport_number,place_of_residence=item.residence)
    try:
        DB.add(passport)
        DB.commit()
        DB.refresh(passport)
        person = Invidivual(first_name=item.first_name,middle_name=item.middle_name,last_name=item.last_name,FK_passports_id=passport.passports_id)
        DB.add(person)
        DB.commit()
        DB.refresh(person)
        return person.individuals_id
    except HTTPException:
        return JSONResponse(status_code=500, content={"message":"Ошибка"})
@alg_router.get("/person_involved/{id}",tags=['Person Involved'])
def get_person_involved(id:int,DB:Session = Depends(get_session)):
    person = DB.query(Invidivual).filter(Invidivual.individuals_id==id).first()
    if person is None:
        return JSONResponse(status_code=404, content={"message":"Человек не найден!"})
    return person
@alg_router.put("/person_involved/edit/{idPerson}",tags=['Person Involved'])
def edit_person_involved_by_id_person_involved(idPerson:int,item:updatePersonInvolved,DB:Session = Depends(get_session)):
    person_involved = DB.query(Invidivual).filter(Invidivual.individuals_id==idPerson).first()
    if person_involved is None:
        return JSONResponse(status_code=404,content={"message":"Участвующее лицо не найдено!"})
    passport = DB.query(Passport).filter(Passport.passports_id == person_involved.FK_passports_id).first()
    if passport is not None:
        passport_unique = DB.query(Passport).filter(item.passport_serial == Passport.passport_serial,item.passport_number== Passport.passport_number,Passport.passports_id != passport.passports_id).first()
        if passport_unique is not None:
            return JSONResponse(status_code=400,content={"message":"Данный паспорт зарегистрирован в системе!"})
    passport = DB.query(Passport).filter(Passport.passports_id == person_involved.FK_passports_id).first()
    person_involved.first_name = item.first_name
    person_involved.middle_name = item.middle_name
    person_involved.last_name = item.last_name
    passport.place_of_residence = item.residence
    passport.passport_serial = item.passport_serial
    passport.passport_number = item.passport_number
    try:
        DB.commit()
        DB.refresh(person_involved)
        DB.commit()
        DB.refresh(passport)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
@alg_router.put("/person_involved/{idQuestPerson}",tags=['Person Involved'])
def edit_person_involved_by_id_quest_person(idQuestPerson:int,item:updatePersonInvolved,DB:Session = Depends(get_session)):
    questperson = DB.query(QuestPerson).filter(QuestPerson.id_quest_persons == idQuestPerson).first()
    if questperson is None:
        return JSONResponse(status_code=404,content={"message":"Допрашиваемый не найден!"})
    if questperson.fk_procedural_position != 8:
        quest = DB.query(Quest).filter(Quest.fk_id_quest_person == idQuestPerson).first()
        if quest is None:
            return JSONResponse(status_code=404,content={"message":"Протокол допроса не найден!"})
        quest_person_involved = DB.query(QuestPersonInvolved).filter(QuestPersonInvolved.id_quest == quest.id_quest).first()
        if quest_person_involved is None:
            return JSONResponse(status_code=404,content={"message":"Участвующие лица не найдены!"})
        person_involved = DB.query(Invidivual).filter(Invidivual.individuals_id == quest_person_involved.id_persons_involved).first()
    else:
        clarif = DB.query(Clarif).filter(Clarif.fk_id_quest_person == questperson.id_quest_persons).first()
        if clarif is None:
            return JSONResponse(status_code=404,content={"message":"Объяснение не найдено!"})
        clarif_person_involved = DB.query(ClarifPersonInvolved).filter(ClarifPersonInvolved.id_clarif == clarif.id_clarif).first()
        if clarif_person_involved is None:
            return JSONResponse(status_code=404,content={"message":"Участвующие лица не найдены!"})
        person_involved = DB.query(Invidivual).filter(Invidivual.individuals_id == clarif_person_involved.id_persons_involved).first()
    passport = DB.query(Passport).filter(Passport.passports_id == person_involved.FK_passports_id).first()
    passport_unique = DB.query(Passport).filter(item.passport_serial == Passport.passport_serial,item.passport_number== Passport.passport_number,Passport.passports_id != passport.passports_id).first()
    if passport_unique is not None:
         return JSONResponse(status_code=400,content={"message":"Данный паспорт зарегистрирован в системе!"})
    passport = DB.query(Passport).filter(Passport.passports_id == person_involved.FK_passports_id).first()
    person_involved.first_name = item.first_name
    person_involved.middle_name = item.middle_name
    person_involved.last_name = item.last_name
    passport.place_of_residence = item.residence
    passport.passport_serial = item.passport_serial
    passport.passport_number = item.passport_number
    try:
        DB.commit()
        DB.refresh(person_involved)
        DB.commit()
        DB.refresh(passport)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})

@alg_router.post("/person_involved_clarif/",tags=['Clarif'])
def create_person_involved_clarif(item:createInspectPersonInvolved,DB:Session = Depends(get_session)):
    clarifperson = DB.query(ClarifPersonInvolved).filter(ClarifPersonInvolved.id_clarif==item.idInspect,ClarifPersonInvolved.id_persons_involved==item.idPerson).first()
    if clarifperson is not None:
        return JSONResponse(status_code=400, content={"message":"Данные уже имеются"})
    clarifperson = ClarifPersonInvolved(id_persons_involved=item.idPerson,id_clarif=item.idInspect,notes=item.notes,petition=item.petition,fk_procedural_position=item.fk_procedural_position)
    try:
        DB.add(clarifperson)
        DB.commit()
        DB.refresh(clarifperson)
        return JSONResponse(status_code=200, content={"message": "Данные добавлены!"})
    except HTTPException:
        return JSONResponse(status_code=500, content={"message":"Ошибка"})
@alg_router.delete("/person_involved_clarif/{idClarif}/{idPersonInvolved}",tags=['Quest'])
def delete_person_involved_clarif_by_id(idClarif:int,idPersonInvolved:int,DB:Session = Depends(get_session)):
    clarif = DB.query(Clarif).filter(idClarif == Clarif.id_clarif).first()
    if clarif is None:
        return JSONResponse(status_code=404, content={"message":"Данные об объяснении не найдены"})
    person_involved = DB.query(Invidivual).filter(idPersonInvolved == Invidivual.individuals_id).first()
    if person_involved is None:
        return JSONResponse(status_code=404, content={"message":"Данные об учавствующем лице не найдены"})
    person_involved_clarif = DB.query(ClarifPersonInvolved).filter(clarif.id_clarif == ClarifPersonInvolved.id_clarif,person_involved.individuals_id == ClarifPersonInvolved.id_persons_involved).first()
    if person_involved_clarif is None:
        return JSONResponse(status_code=404, content={"message":"Данные об учавствующем лице не найдены"})
    try:
        DB.delete(person_involved_clarif)
        DB.commit()
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
@alg_router.put("/person_involved_clarif/{idClarif}/{idPersonInvolved}",tags=['Quest'])
def edit_person_involved_clarif_by_id(idClarif:int,idPersonInvolved:int,item:updatePersonInvolvedQuest,DB:Session = Depends(get_session)):
    clarif = DB.query(Clarif).filter(idClarif == Clarif.id_clarif).first()
    if clarif is None:
        return JSONResponse(status_code=404, content={"message":"Данные об объяснении не найдены"})
    person_involved = DB.query(Invidivual).filter(idPersonInvolved == Invidivual.individuals_id).first()
    if person_involved is None:
        return JSONResponse(status_code=404, content={"message":"Данные об учавствующем лице не найдены"})
    person_involved_clarif = DB.query(ClarifPersonInvolved).filter(clarif.id_clarif == ClarifPersonInvolved.id_clarif,person_involved.individuals_id == ClarifPersonInvolved.id_persons_involved).first()
    if person_involved_clarif is None:
        return JSONResponse(status_code=404, content={"message":"Данные об учавствующем лице не найдены"})
    person_involved_clarif.notes = item.notes
    person_involved_clarif.petition = item.petition
    try:
        DB.commit()
        DB.refresh(person_involved_clarif)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
@alg_router.get("/person_involved_clarif/{idClarifPerson}",tags=['Clarif'])
def get_all_person_clarif(idClarifPerson:int,DB:Session = Depends(get_session)):
    clarif= DB.query(Clarif).filter(idClarifPerson == Clarif.fk_id_quest_person).first()
    if clarif is None:
        return JSONResponse(status_code=404, content={"message":"Данные об объяснении не найдены"})
    clarifperson = DB.query(ClarifPersonInvolved).filter(ClarifPersonInvolved.id_clarif==clarif.id_clarif).all()
    if len(clarifperson) == 0:
        return JSONResponse(status_code=404,content={"message":"Данные не найдены"})
    all_person = []
    for i in range(len(clarifperson)):
        person = DB.query(Invidivual).filter(Invidivual.individuals_id == clarifperson[i].id_persons_involved).first()
        passport = DB.query(Passport).filter(person.FK_passports_id==Passport.passports_id).first()
        if clarifperson[i].petition == 1:
            petition = 'Нет'
        else:
            petition = 'Да'
        procedural_position_name = DB.query(ProceduralPosition).filter(clarifperson[i].fk_procedural_position==ProceduralPosition.id_procedural_position).first().name_position
        info = {
            "id_persons_involved":person.individuals_id,
            "first_name":person.first_name,
            "middle_name":person.middle_name,
            "last_name":person.last_name,
            "passport_serial":passport.passport_serial,
            "passport_number":passport.passport_number,
            "residence":passport.place_of_residence,
            "fk_procedural_position":procedural_position_name,
            "notes":clarifperson[i].notes,
            "petition":petition
        }
        all_person.append(info)
    return all_person
@alg_router.post("/clarif/",tags=['Clarif'])
def clarif_create(item:createClarif,DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == item.user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == item.case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    if cases.fk_user_id != user.id:
        return JSONResponse(status_code= 400, content={"message": "Этот пользователь не может редактировать данное дело"})
    clarif = DB.query(Clarif).filter(Clarif.fk_id_quest_person == item.quest_person).first()
    if clarif is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    clarif = Clarif(date_clarif=item.date_clarif,begin_clarif=item.begin_clarif,end_clarif=item.end_clarif,place_clarif=item.place_clarif,text_clarif=item.text_clarif,technical_means=item.technical_means,read_clarif=item.read_clarif,true_clarif=item.true_clarif,addition_clarif=item.addition_clarif,photography=item.photography,fk_id_quest_person=item.quest_person)
    if clarif is None:
        return JSONResponse(status_code= 404, content={"message": "Произошла ошибка!"})
    try:
        DB.add(clarif)
        DB.commit()
        DB.refresh(clarif)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/clarif/{QuestPersonId}",tags=['Clarif'])
def get_clarif(QuestPersonId:int,DB:Session = Depends(get_session)):
    questperson = DB.query(QuestPerson).filter(QuestPerson.id_quest_persons == QuestPersonId).first()
    if questperson is None:
        return JSONResponse(status_code= 404, content={"message": "Заявитель не найден"})
    caseId = questperson.fk_active_case_id
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    clarif = DB.query(Clarif).filter(Clarif.fk_id_quest_person==QuestPersonId).first()
    if clarif is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return clarif
@alg_router.put("/clarif/{QuestPersonId}",tags=['Clarif'])
def edit_clarif(QuestPersonId:int,item:updateClarif,DB:Session = Depends(get_session)):
    questperson = DB.query(QuestPerson).filter(QuestPerson.id_quest_persons == QuestPersonId).first()
    if questperson is None:
        return JSONResponse(status_code= 404, content={"message": "Заявитель не найден"})
    caseId = questperson.fk_active_case_id
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    clarif = DB.query(Clarif).filter(Clarif.fk_id_quest_person==QuestPersonId).first()
    if clarif is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    clarif.date_clarif = item.date_clarif
    clarif.begin_clarif = item.begin_clarif
    clarif.end_clarif = item.end_clarif
    clarif.place_clarif = item.place_clarif
    clarif.text_clarif = item.text_clarif
    clarif.technical_means = item.technical_means
    clarif.read_clarif = item.read_clarif
    clarif.true_clarif = item.true_clarif
    clarif.addition_clarif = item.addition_clarif
    clarif.photography = item.photography
    try:
        DB.commit()
        DB.refresh(clarif)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
@alg_router.post("/track/",tags=['Track'])
def track_create(item:createTrack,DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == item.user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == item.case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    if cases.fk_user_id != user.id:
        return JSONResponse(status_code= 400, content={"message": "Этот пользователь не может редактировать данное дело"})
    track = DB.query(Track).filter(Track.fk_active_case_id == item.case_id,Track.type_track==item.type_track,Track.info_track == item.info_track).first()
    if track is not None:
        return JSONResponse(status_code= 400, content={"message": "Цифровой след уже добавлен!"})
    track = Track(fk_active_case_id=item.case_id,type_track = item.type_track,info_track = item.info_track)
    if track is None:
        return JSONResponse(status_code= 404, content={"message": "Произошла ошибка!"})
    try:
        DB.add(track)
        DB.commit()
        DB.refresh(track)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/track/{caseId}",tags=['Track'])
def get_track(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    track = DB.query(Track).filter(Track.fk_active_case_id==caseId).order_by(Track.id_track).all()
    if len(track) == 0:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return track
@alg_router.put("/track/{track_id}",tags=['Track'])
def edit_track(track_id:int,item:updateTrack,DB:Session = Depends(get_session)):
    track = DB.query(Track).filter(Track.id_track==track_id).first()
    if track is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    cases = DB.query(Case).filter(Case.active_cases_id == track.fk_active_case_id).first()
    track = DB.query(Track).filter(Track.fk_active_case_id==cases.active_cases_id,Track.info_track==item.info_track,Track.type_track == item.type_track).first()
    if track is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже имеются"})
    track = DB.query(Track).filter(Track.id_track==track_id).first()
    track.type_track = item.type_track
    track.info_track = item.info_track
    try:
        DB.commit()
        DB.refresh(track)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
@alg_router.delete("/track/{track_id}",tags=['Track'])
def delete_track(track_id:int,DB:Session = Depends(get_session)):
    track = DB.query(Track).filter(Track.id_track==track_id).first()
    if track is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    try:
        DB.delete(track)
        DB.commit()
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
@alg_router.post("/person_involved_inspect/",tags=['Inspect'])
def create_inspect_person_involved(item:createInspectPersonInvolved,DB:Session = Depends(get_session)):
    inspectperson = DB.query(InspectPersonsInvolved).filter(InspectPersonsInvolved.id_inspect==item.idInspect,InspectPersonsInvolved.id_persons_involved==item.idPerson).first()
    if inspectperson is not None:
        return JSONResponse(status_code=400, content={"message":"Данные уже имеются"})
    inspectperson = InspectPersonsInvolved(id_persons_involved=item.idPerson,id_inspect=item.idInspect,notes=item.notes,petition=item.petition,fk_procedural_position=item.fk_procedural_position)
    try:
        DB.add(inspectperson)
        DB.commit()
        DB.refresh(inspectperson)
        return JSONResponse(status_code=200, content={"message": "Данные добавлены!"})
    except HTTPException:
        return JSONResponse(status_code=500, content={"message":"Ошибка"})
@alg_router.put("/person_involved_inspect/{idInspect}/{idPersonInvolved}",tags=['Inspect'])
def edit_person_involved_inspect_by_id(idInspect:int,idPersonInvolved:int,item:updatePersonInvolvedQuest,DB:Session = Depends(get_session)):
    inspect = DB.query(Inspect).filter(idInspect == Inspect.id_inspect).first()
    if inspect is None:
        return JSONResponse(status_code=404, content={"message":"Данные о протоколе не найдены"})
    person_involved = DB.query(Invidivual).filter(idPersonInvolved == Invidivual.individuals_id).first()
    if person_involved is None:
        return JSONResponse(status_code=404, content={"message":"Данные об учавствующем лице не найдены"})
    person_involved_inspect = DB.query(InspectPersonsInvolved).filter(inspect.id_inspect == InspectPersonsInvolved.id_inspect,person_involved.individuals_id == InspectPersonsInvolved.id_persons_involved).first()
    if person_involved_inspect is None:
        return JSONResponse(status_code=404, content={"message":"Данные об учавствующем лице не найдены"})
    person_involved_inspect.notes = item.notes
    person_involved_inspect.petition = item.petition
    person_involved_inspect.fk_procedural_position = item.fk_procedural_position
    try:
        DB.commit()
        DB.refresh(person_involved_inspect)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
@alg_router.delete("/person_involved_inspect/{idInspect}/{idPersonInvolved}",tags=['Inspect'])
def delete_person_involved_inspect_by_id(idInspect:int,idPersonInvolved:int,DB:Session = Depends(get_session)):
    inspect = DB.query(Inspect).filter(idInspect == Inspect.id_inspect).first()
    if inspect is None:
        return JSONResponse(status_code=404, content={"message":"Данные о протоколе не найдены"})
    person_involved = DB.query(Invidivual).filter(idPersonInvolved == Invidivual.individuals_id).first()
    if person_involved is None:
        return JSONResponse(status_code=404, content={"message":"Данные об учавствующем лице не найдены"})
    person_involved_inspect = DB.query(InspectPersonsInvolved).filter(inspect.id_inspect == InspectPersonsInvolved.id_inspect,person_involved.individuals_id == InspectPersonsInvolved.id_persons_involved).first()
    if person_involved_inspect is None:
        return JSONResponse(status_code=404, content={"message":"Данные об учавствующем лице не найдены"})
    try:
        DB.delete(person_involved_inspect)
        DB.commit()
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
@alg_router.get("/person_involved_inspect/{id}",tags=['Inspect'])
def get_all_inspects_person(id:int,DB:Session = Depends(get_session)):
    inspect_id = DB.query(Inspect).filter(id == Inspect.fk_active_case_id).first().id_inspect
    inspectperson = DB.query(InspectPersonsInvolved).filter(InspectPersonsInvolved.id_inspect==inspect_id).order_by(InspectPersonsInvolved.id_persons_involved).all()
    if len(inspectperson) == 0:
        return JSONResponse(status_code=404,content={"message":"Данные не найдены"})
    all_person = []
    for i in range(len(inspectperson)):
        person = DB.query(Invidivual).filter(Invidivual.individuals_id == inspectperson[i].id_persons_involved).first()
        passport = DB.query(Passport).filter(person.FK_passports_id == Passport.passports_id).first()
        if inspectperson[i].petition == 1:
            petition = 'Нет'
        else:
            petition = 'Да'
        procedural_position_name = DB.query(ProceduralPosition).filter(inspectperson[i].fk_procedural_position==ProceduralPosition.id_procedural_position).first().name_position
        info = {
            "id_persons_involved":person.individuals_id,
            "first_name":person.first_name,
            "middle_name":person.middle_name,
            "last_name":person.last_name,
            "passport_serial":passport.passport_serial,
            "passport_number":passport.passport_number,
            "residence":passport.place_of_residence,
            "fk_procedural_position":procedural_position_name,
            "notes":inspectperson[i].notes,
            "petition":petition
        }
        all_person.append(info)
    return all_person
@alg_router.post("/inspect/",tags=['Inspect'])
def inspect_create(item:createInspect,DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == item.user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == item.case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    if cases.fk_user_id != user.id:
        return JSONResponse(status_code= 400, content={"message": "Этот пользователь не может редактировать данное дело"})
    inspect = DB.query(Inspect).filter(Inspect.fk_active_case_id == item.case_id).first()
    if inspect is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    inspect = Inspect(fk_active_case_id=item.case_id,date_inspect=item.date_inspect,begin_inspect=item.begin_inspect,end_inspect=item.end_inspect,message_inspect=item.message_inspect,from_message_inspect=item.from_message_inspect,place_of_inspect=item.place_of_inspect,technical_means=item.technical_means,conditions=item.conditions,establish=item.establish,photography=item.photography,sized_items=item.sized_items,items_for_inspect=item.items_for_inspect,familiarization=item.familiarization,inspect_exam=item.inspect_exam)
    if inspect is None:
        return JSONResponse(status_code= 404, content={"message": "Произошла ошибка!"})
    try:
        DB.add(inspect)
        DB.commit()
        DB.refresh(inspect)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/inspect/{caseId}",tags=['Inspect'])
def get_inspect(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    inspect = DB.query(Inspect).filter(Inspect.fk_active_case_id==caseId).first()
    if inspect is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return inspect
@alg_router.put("/inspect/{caseId}",tags=['Inspect'])
def edit_inspect(caseId:int,item:updateInspect,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    inspect = DB.query(Inspect).filter(Inspect.fk_active_case_id==caseId).first()
    if inspect is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    inspect.date_inspect = item.date_inspect
    inspect.begin_inspect = item.begin_inspect
    inspect.end_inspect=item.end_inspect
    inspect.message_inspect=item.message_inspect
    inspect.from_message_inspect=item.from_message_inspect
    inspect.place_of_inspect=item.place_of_inspect
    inspect.technical_means=item.technical_means
    inspect.conditions=item.conditions
    inspect.establish=item.establish
    inspect.photography=item.photography
    inspect.sized_items=item.sized_items
    inspect.items_for_inspect=item.items_for_inspect
    inspect.familiarization=item.familiarization
    inspect.inspect_exam=item.inspect_exam
    try:
        DB.commit()
        DB.refresh(inspect)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
@alg_router.post("/reclaim/{caseId}/upload_passport",tags=['Reclaim'])
def reclaim_upload_passport(caseId:int,copyes_passport:List[UploadFile] = File(...),DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}') == False:
        return JSONResponse(status_code= 404, content={"message": "Добавление файлов невозможно!"})
    for copy in copyes_passport:
        copy.filename = copy.filename.lower()
        path = f'media/{id}/cases/{caseId}/reclaim/passport/{copy.filename}'
        with open(path,'wb+') as buffer:
            shutil.copyfileobj(copy.file,buffer)
    return JSONResponse(status_code= 200, content={"message": "Файлы успешно загружены"})
@alg_router.post("/reclaim/{caseId}/upload_income",tags=['Reclaim'])
def reclaim_upload_income(caseId:int,copyes_passport:List[UploadFile] = File(...),DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}') == False:
        return JSONResponse(status_code= 404, content={"message": "Добавление файлов невозможно!"})
    for copy in copyes_passport:
        copy.filename = copy.filename.lower()
        path = f'media/{id}/cases/{caseId}/reclaim/income/{copy.filename}'
        with open(path,'wb+') as buffer:
            shutil.copyfileobj(copy.file,buffer)
    return JSONResponse(status_code= 200, content={"message": "Файлы успешно загружены"})
@alg_router.post("/reclaim/{caseId}/upload_detailing",tags=['Reclaim'])
def reclaim_upload_detailing(caseId:int,copyes_passport:List[UploadFile] = File(...),DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}') == False:
        return JSONResponse(status_code= 404, content={"message": "Добавление файлов невозможно!"})
    for copy in copyes_passport:
        copy.filename = copy.filename.lower()
        path = f'media/{id}/cases/{caseId}/reclaim/detailing/{copy.filename}'
        with open(path,'wb+') as buffer:
            shutil.copyfileobj(copy.file,buffer)
    return JSONResponse(status_code= 200, content={"message": "Файлы успешно загружены"})
@alg_router.get("/reclaim/{caseId}/get_passport",tags=['Reclaim'])
def get_passports_lists_filenames(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}') == False:
        return JSONResponse(status_code= 404, content={"message": "Добавление файлов невозможно!"})
    files = os.listdir(f'media/{id}/cases/{caseId}/reclaim/passport')
    if len(files) == 0:
        return JSONResponse(status_code=404,content={"message":"Файлы не найдены"})
    return files
@alg_router.get("/reclaim/{caseId}/get_income",tags=['Reclaim'])
def get_income_lists_filenames(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}') == False:
        return JSONResponse(status_code= 404, content={"message": "Добавление файлов невозможно!"})
    files = os.listdir(f'media/{id}/cases/{caseId}/reclaim/income')
    if len(files) == 0:
        return JSONResponse(status_code=404,content={"message":"Файлы не найдены"})
    return files
@alg_router.get("/reclaim/{caseId}/get_detailing",tags=['Reclaim'])
def get_detailing_lists_filenames(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}') == False:
        return JSONResponse(status_code= 404, content={"message": "Добавление файлов невозможно!"})
    files = os.listdir(f'media/{id}/cases/{caseId}/reclaim/detailing')
    if len(files) == 0:
        return JSONResponse(status_code=404,content={"message":"Файлы не найдены"})
    return files
@alg_router.get("/reclaim/{caseId}/passport_download/{filenames}",tags=['Reclaim'])
def download_passport_file(caseId:int,filenames:str,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}/reclaim/passport/{filenames}') == False:
        return JSONResponse(status_code= 404, content={"message": "Файл не существует"})
    return FileResponse(f'media/{id}/cases/{caseId}/reclaim/passport/{filenames}',filename=f'{filenames}',media_type='application/octet-stream')
@alg_router.get("/reclaim/{caseId}/income_download/{filenames}",tags=['Reclaim'])
def download_income_file(caseId:int,filenames:str,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}/reclaim/income/{filenames}') == False:
        return JSONResponse(status_code= 404, content={"message": "Файл не существует"})
    return FileResponse(f'media/{id}/cases/{caseId}/reclaim/income/{filenames}',filename=f'{filenames}',media_type='application/octet-stream')
@alg_router.get("/reclaim/{caseId}/detailing_download/{filenames}",tags=['Reclaim'])
def download_detailing_file(caseId:int,filenames:str,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}/reclaim/detailing/{filenames}') == False:
        return JSONResponse(status_code= 404, content={"message": "Файл не существует"})
    return FileResponse(f'media/{id}/cases/{caseId}/reclaim/detailing/{filenames}',filename=f'{filenames}',media_type='application/octet-stream')
@alg_router.delete("/reclaim/{caseId}/passport_delete/{filenames}",tags=['Reclaim'])
def delete_passport_file(caseId:int,filenames:str,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}/reclaim/passport/{filenames}') == False:
        return JSONResponse(status_code= 404, content={"message": "Файл не существует"})
    os.remove(f'media/{id}/cases/{caseId}/reclaim/passport/{filenames}')
    return JSONResponse(status_code= 200, content={"message": "Файл удален!"})
@alg_router.delete("/reclaim/{caseId}/income_delete/{filenames}",tags=['Reclaim'])
def delete_income_file(caseId:int,filenames:str,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}/reclaim/income/{filenames}') == False:
        return JSONResponse(status_code= 404, content={"message": "Файл не существует"})
    os.remove(f'media/{id}/cases/{caseId}/reclaim/income/{filenames}')
    return JSONResponse(status_code= 200, content={"message": "Файл удален!"})
@alg_router.delete("/reclaim/{caseId}/detailing_delete/{filenames}",tags=['Reclaim'])
def delete_detailing_file(caseId:int,filenames:str,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}/reclaim/detailing/{filenames}') == False:
        return JSONResponse(status_code= 404, content={"message": "Файл не существует"})
    os.remove(f'media/{id}/cases/{caseId}/reclaim/detailing/{filenames}')
    return JSONResponse(status_code= 200, content={"message": "Файл удален!"})

@alg_router.post("/extract/{caseId}/upload_money_orders",tags=['Extract'])
def extract_upload_money_orders(caseId:int,copyes_passport:List[UploadFile] = File(...),DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}') == False:
        return JSONResponse(status_code= 404, content={"message": "Добавление файлов невозможно!"})
    for copy in copyes_passport:
        copy.filename = copy.filename.lower()
        path = f'media/{id}/cases/{caseId}/extract/money_orders/{copy.filename}'
        with open(path,'wb+') as buffer:
            shutil.copyfileobj(copy.file,buffer)
    return JSONResponse(status_code= 200, content={"message": "Файлы успешно загружены"})
@alg_router.post("/extract/{caseId}/upload_detailing",tags=['Extract'])
def extract_upload_detailing(caseId:int,copyes_passport:List[UploadFile] = File(...),DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}') == False:
        return JSONResponse(status_code= 404, content={"message": "Добавление файлов невозможно!"})
    for copy in copyes_passport:
        copy.filename = copy.filename.lower()
        path = f'media/{id}/cases/{caseId}/extract/detailing/{copy.filename}'
        with open(path,'wb+') as buffer:
            shutil.copyfileobj(copy.file,buffer)
    return JSONResponse(status_code= 200, content={"message": "Файлы успешно загружены"})
@alg_router.post("/extract/{caseId}/upload_screenshot",tags=['Extract'])
def extract_upload_screenshot(caseId:int,copyes_passport:List[UploadFile] = File(...),DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}') == False:
        return JSONResponse(status_code= 404, content={"message": "Добавление файлов невозможно!"})
    for copy in copyes_passport:
        copy.filename = copy.filename.lower()
        path = f'media/{id}/cases/{caseId}/extract/screenshot/{copy.filename}'
        with open(path,'wb+') as buffer:
            shutil.copyfileobj(copy.file,buffer)
    return JSONResponse(status_code= 200, content={"message": "Файлы успешно загружены"})
@alg_router.get("/extract/{caseId}/get_detailing",tags=['Extract'])
def get_detailing_lists_filenames(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}') == False:
        return JSONResponse(status_code= 404, content={"message": "Добавление файлов невозможно!"})
    files = os.listdir(f'media/{id}/cases/{caseId}/extract/detailing')
    if len(files) == 0:
        return JSONResponse(status_code=404,content={"message":"Файлы не найдены"})
    return files
@alg_router.get("/extract/{caseId}/get_screenshot",tags=['Extract'])
def get_screenshot_lists_filenames(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}') == False:
        return JSONResponse(status_code= 404, content={"message": "Добавление файлов невозможно!"})
    files = os.listdir(f'media/{id}/cases/{caseId}/extract/screenshot')
    if len(files) == 0:
        return JSONResponse(status_code=404,content={"message":"Файлы не найдены"})
    return files
@alg_router.get("/extract/{caseId}/get_money_orders",tags=['Extract'])
def get_money_order_lists_filenames(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}') == False:
        return JSONResponse(status_code= 404, content={"message": "Добавление файлов невозможно!"})
    files = os.listdir(f'media/{id}/cases/{caseId}/extract/money_orders')
    if len(files) == 0:
        return JSONResponse(status_code=404,content={"message":"Файлы не найдены"})
    return files
@alg_router.get("/extract/{caseId}/detailing_download/{filenames}",tags=['Extract'])
def download_detailing_file(caseId:int,filenames:str,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}/extract/detailing/{filenames}') == False:
        return JSONResponse(status_code= 404, content={"message": "Файл не существует"})
    return FileResponse(f'media/{id}/cases/{caseId}/extract/detailing/{filenames}',filename=f'{filenames}',media_type='application/octet-stream')
@alg_router.get("/extract/{caseId}/screenshot_download/{filenames}",tags=['Extract'])
def download_screenshot_file(caseId:int,filenames:str,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}/extract/screenshot/{filenames}') == False:
        return JSONResponse(status_code= 404, content={"message": "Файл не существует"})
    return FileResponse(f'media/{id}/cases/{caseId}/extract/screenshot/{filenames}',filename=f'{filenames}',media_type='application/octet-stream')
@alg_router.get("/extract/{caseId}/money_orders_download/{filenames}",tags=['Extract'])
def download_money_orders_file(caseId:int,filenames:str,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}/extract/money_orders/{filenames}') == False:
        return JSONResponse(status_code= 404, content={"message": "Файл не существует"})
    return FileResponse(f'media/{id}/cases/{caseId}/extract/money_orders/{filenames}',filename=f'{filenames}',media_type='application/octet-stream')
@alg_router.delete("/extract/{caseId}/detailing_delete/{filenames}",tags=['Extract'])
def delete_detailing_file(caseId:int,filenames:str,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}/extract/detailing/{filenames}') == False:
        return JSONResponse(status_code= 404, content={"message": "Файл не существует"})
    os.remove(f'media/{id}/cases/{caseId}/extract/detailing/{filenames}')
    return JSONResponse(status_code= 200, content={"message": "Файл удален!"})
@alg_router.delete("/extract/{caseId}/money_orders_delete/{filenames}",tags=['Extract'])
def delete_money_orders_file(caseId:int,filenames:str,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}/extract/money_orders/{filenames}') == False:
        return JSONResponse(status_code= 404, content={"message": "Файл не существует"})
    os.remove(f'media/{id}/cases/{caseId}/extract/money_orders/{filenames}')
    return JSONResponse(status_code= 200, content={"message": "Файл удален!"})
@alg_router.delete("/extract/{caseId}/screenshot_delete/{filenames}",tags=['Extract'])
def delete_screenshot_file(caseId:int,filenames:str,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    id = user.id
    if os.path.exists(f'media/{id}/cases/{caseId}/extract/screenshot/{filenames}') == False:
        return JSONResponse(status_code= 404, content={"message": "Файл не существует"})
    os.remove(f'media/{id}/cases/{caseId}/extract/screenshot/{filenames}')
    return JSONResponse(status_code= 200, content={"message": "Файл удален!"})
@alg_router.post("/person_involved_check/",tags=['Check'])
def create_person_involved_check(item:createInspectPersonInvolved,DB:Session = Depends(get_session)):
    checkperson = DB.query(CheckPersonsInvolved).filter(CheckPersonsInvolved.id_check==item.idInspect,CheckPersonsInvolved.id_persons_involved==item.idPerson).first()
    if checkperson is not None:
        return JSONResponse(status_code=400, content={"message":"Данные уже имеются"})
    checkperson = CheckPersonsInvolved(id_persons_involved=item.idPerson,id_check=item.idInspect,notes=item.notes,petition=item.petition,fk_procedural_position=item.fk_procedural_position)
    try:
        DB.add(checkperson)
        DB.commit()
        DB.refresh(checkperson)
        return JSONResponse(status_code=200, content={"message": "Данные добавлены!"})
    except HTTPException:
        return JSONResponse(status_code=500, content={"message":"Ошибка"})
@alg_router.put("/person_involved_check/{idCheck}/{idPersonInvolved}",tags=['Check'])
def edit_person_involved_check_by_id(idCheck:int,idPersonInvolved:int,item:updatePersonInvolvedQuest,DB:Session = Depends(get_session)):
    check = DB.query(Check).filter(idCheck == Check.id_check).first()
    if check is None:
        return JSONResponse(status_code=404, content={"message":"Данные о протоколе не найдены"})
    person_involved = DB.query(Invidivual).filter(idPersonInvolved == Invidivual.individuals_id).first()
    if person_involved is None:
        return JSONResponse(status_code=404, content={"message":"Данные об учавствующем лице не найдены"})
    person_involved_check = DB.query(CheckPersonsInvolved).filter(check.id_check == CheckPersonsInvolved.id_check,person_involved.individuals_id == CheckPersonsInvolved.id_persons_involved).first()
    if person_involved_check is None:
        return JSONResponse(status_code=404, content={"message":"Данные об учавствующем лице не найдены"})
    person_involved_check.notes = item.notes
    person_involved_check.petition = item.petition
    person_involved_check.fk_procedural_position = item.fk_procedural_position
    try:
        DB.commit()
        DB.refresh(person_involved_check)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
@alg_router.delete("/person_involved_check/{idCheck}/{idPersonInvolved}",tags=['Check'])
def delete_person_involved_check_by_id(idCheck:int,idPersonInvolved:int,DB:Session = Depends(get_session)):
    check = DB.query(Check).filter(idCheck == Check.id_check).first()
    if check is None:
        return JSONResponse(status_code=404, content={"message":"Данные о протоколе не найдены"})
    person_involved = DB.query(Invidivual).filter(idPersonInvolved == Invidivual.individuals_id).first()
    if person_involved is None:
        return JSONResponse(status_code=404, content={"message":"Данные об учавствующем лице не найдены"})
    person_involved_check = DB.query(CheckPersonsInvolved).filter(check.id_check == CheckPersonsInvolved.id_check,person_involved.individuals_id == CheckPersonsInvolved.id_persons_involved).first()
    if person_involved_check is None:
        return JSONResponse(status_code=404, content={"message":"Данные об учавствующем лице не найдены"})
    try:
        DB.delete(person_involved_check)
        DB.commit()
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
@alg_router.get("/person_involved_check/{id}",tags=['Check'])
def get_all_person_check(id:int,DB:Session = Depends(get_session)):
    check_id = DB.query(Check).filter(id == Check.fk_active_case_id).first().id_check
    checkperson = DB.query(CheckPersonsInvolved).filter(CheckPersonsInvolved.id_check==check_id).all()
    if len(checkperson) == 0:
        return JSONResponse(status_code=404,content={"message":"Данные не найдены"})
    all_person = []
    for i in range(len(checkperson)):
        person = DB.query(Invidivual).filter(Invidivual.individuals_id == checkperson[i].id_persons_involved).first()
        passport = DB.query(Passport).filter(Passport.passports_id == person.FK_passports_id).first()
        if checkperson[i].petition == 1:
            petition = 'Нет'
        else:
            petition = 'Да'
        procedural_position_name = DB.query(ProceduralPosition).filter(checkperson[i].fk_procedural_position==ProceduralPosition.id_procedural_position).first().name_position
        info = {
            "id_persons_involved":person.individuals_id,
            "first_name":person.first_name,
            "middle_name":person.middle_name,
            "last_name":person.last_name,
            "passport_serial":passport.passport_serial,
            "passport_number":passport.passport_number,
            "residence":passport.place_of_residence,
            "fk_procedural_position":procedural_position_name,
            "notes":checkperson[i].notes,
            "petition":petition
        }
        all_person.append(info)
    return all_person
@alg_router.post("/check/",tags=['Check'])
def check_create(item:createCheck,DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == item.user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == item.case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    if cases.fk_user_id != user.id:
        return JSONResponse(status_code= 400, content={"message": "Этот пользователь не может редактировать данное дело"})
    check = DB.query(Check).filter(Check.fk_active_case_id == item.case_id).first()
    if check is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    check = Check(fk_active_case_id=item.case_id,items=item.items,date_check=item.date_check,begin_check=item.begin_check,end_check=item.end_check,check_what=item.check_what,check_place=item.check_place,when_size=item.when_size,where_size=item.where_size,which_action_size=item.which_action_size,technical_means=item.technical_means,conditions=item.conditions,establish=item.establish,wrap=item.wrap,photography=item.photography,items_for_check=item.items_for_check,familiarizathion=item.familiarizathion)
    if check is None:
        return JSONResponse(status_code= 404, content={"message": "Произошла ошибка!"})
    try:
        DB.add(check)
        DB.commit()
        DB.refresh(check)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/check/{caseId}",tags=['Check'])
def get_check(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    check = DB.query(Check).filter(Check.fk_active_case_id==caseId).first()
    if check is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return check
@alg_router.put("/check/{caseId}",tags=['Check'])
def edit_check(caseId:int,item:updateCheck,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    check = DB.query(Check).filter(Check.fk_active_case_id==caseId).first()
    if check is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    check.items = item.items
    check.date_check = item.date_check
    check.begin_check = item.begin_check
    check.end_check = item.end_check
    check.check_what = item.check_what
    check.check_place = item.check_place
    check.when_size = item.when_size
    check.where_size = item.where_size
    check.which_action_size = item.which_action_size
    check.technical_means = item.technical_means
    check.conditions = item.conditions
    check.establish = item.establish
    check.wrap = item.wrap
    check.photography = item.photography
    check.items_for_check = item.items_for_check
    check.familiarizathion = item.familiarizathion
    try:
        DB.commit()
        DB.refresh(check)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})

@alg_router.post("/expertise/",tags=['Expertise'])
def expertise_create(item:createExpertise,DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == item.user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == item.case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    if cases.fk_user_id != user.id:
        return JSONResponse(status_code= 400, content={"message": "Этот пользователь не может редактировать данное дело"})
    expertise = DB.query(Expertise).filter(Expertise.fk_active_case_id == item.case_id).first()
    if expertise is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    expertise = Expertise(fk_active_case_id=item.case_id,app_comp=item.app_comp,prog_comp=item.prog_comp,inf_comp=item.inf_comp,comp_net=item.comp_net)
    if expertise is None:
        return JSONResponse(status_code= 404, content={"message": "Произошла ошибка!"})
    try:
        DB.add(expertise)
        DB.commit()
        DB.refresh(expertise)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/expertise/{caseId}",tags=['Expertise'])
def get_expertise(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    expertise = DB.query(Expertise).filter(Expertise.fk_active_case_id==caseId).first()
    if expertise is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return expertise
@alg_router.put("/expertise/{caseId}",tags=['Expertise'])
def edit_expertise(caseId:int,item:updateExpertise,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    expertise = DB.query(Expertise).filter(Expertise.fk_active_case_id==caseId).first()
    if expertise is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    expertise.app_comp = item.app_comp
    expertise.prog_comp = item.prog_comp
    expertise.inf_comp = item.inf_comp
    expertise.comp_net = item.comp_net
    try:
        DB.commit()
        DB.refresh(expertise)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})

@alg_router.post("/quest_person/create_victim",tags=['Quest Person'])
def create_victim_quest_person(item:createVictimQuestPerson,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == item.fk_active_case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message":"Дело не найдено"})
    if (item.fk_procedural_position == 1):
        declaration_victim = DB.query(DeclarationVictim).filter(DeclarationVictim.fk_active_case_id == cases.active_cases_id).first()
        if declaration_victim is None:
            return JSONResponse(status_code= 404, content={"message":"Постановление о признании потерпевшим не найдено!"})
        id_individ = declaration_victim.fk_individuals_id
    else:
        statement = DB.query(Statement).filter(Statement.statement_id == cases.FK_statement_id).first()
        if statement is None:
            return JSONResponse(status_code= 404, content={"message":"Заявление не найдено!"})
        id_individ = statement.fk_individuals_id
    quest_person = QuestPerson(fk_active_case_id=item.fk_active_case_id,fk_procedural_position=item.fk_procedural_position,individuals_id=id_individ,notes=item.notes,petition = item.petition)
    try:
        DB.add(quest_person)
        DB.commit()
        DB.refresh(quest_person)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.post("/quest_person",tags=['Quest Person'])
def create_quest_person(item:createQuestPerson,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == item.fk_active_case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message":"Дело не найдено"})
    passport = DB.query(Passport).filter(Passport.passport_serial == item.passport_serial, Passport.passport_number == item.passport_number).first()
    if passport is not None:
        individual = DB.query(Invidivual).filter(passport.passports_id == Invidivual.FK_passports_id,Invidivual.first_name == item.firstname,Invidivual.middle_name == item.middlename, Invidivual.last_name == item.lastname).first()
        if individual is None:
            return JSONResponse(status_code= 404, content={"message":"Данный паспорт уже зарегистрирован в системе"})
        individual.resident=item.resident
        individual.education=item.education
        individual.family_status=item.family_status
        individual.place_work=item.place_work
        individual.military=item.military
        individual.criminal=item.criminal
        individual.other=item.other
        DB.commit()
        DB.refresh(passport)
    else:
        passport = Passport(passport_serial=item.passport_serial,passport_number=item.passport_number,birthdate=item.datebrithday,place_of_birth=item.bitrh_residence,place_of_residence=item.place_of_recidense)
        DB.add(passport)
        DB.commit()
        DB.refresh(passport)
        individual = Invidivual(first_name=item.firstname,middle_name=item.middlename,last_name=item.lastname,FK_passports_id=passport.passports_id,resident=item.resident,education=item.education,family_status=item.family_status,place_work=item.place_work,military=item.military,criminal=item.criminal,other = item.other)
        DB.add(individual)
        DB.commit()
        DB.refresh(individual)
    questperson = DB.query(QuestPerson).filter(QuestPerson.fk_active_case_id==cases.active_cases_id,QuestPerson.individuals_id==individual.individuals_id,QuestPerson.fk_procedural_position == item.fk_procedural_position).first()
    if questperson is not None:
        return JSONResponse(status_code= 400, content={"message":"Данные уже имеются"})
    questperson = QuestPerson(fk_active_case_id=item.fk_active_case_id,fk_procedural_position=item.fk_procedural_position,individuals_id=individual.individuals_id,notes=item.notes,petition = item.petition)
    try:
        DB.add(questperson)
        DB.commit()
        DB.refresh(questperson)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.delete("/quest_person/{idQuestPerson}",tags=['Quest Person'])
def delete_quest_person(idQuestPerson:int,DB:Session = Depends(get_session)):
    questperson = DB.query(QuestPerson).filter(QuestPerson.id_quest_persons == idQuestPerson).first()
    if questperson is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    quest = DB.query(Quest).filter(Quest.fk_id_quest_person == questperson.id_quest_persons).first()
    if quest is None:
        DB.delete(questperson)
        DB.commit()
        return JSONResponse(status_code=200,content={"message":"Данные успешно удалены!"})
    quest_person_involved = DB.query(QuestPersonInvolved).filter(QuestPersonInvolved.id_quest == quest.id_quest).first()
    if quest_person_involved is None:
        DB.delete(quest)
        DB.commit()
        DB.delete(questperson)
        DB.commit()
        return JSONResponse(status_code=200,content={"message":"Данные успешно удалены!"})
    DB.delete(quest_person_involved)
    DB.commit()
    DB.delete(quest)
    DB.commit()
    DB.delete(questperson)
    DB.commit()
    return JSONResponse(status_code=200,content={"message":"Данные успешно удалены!"})

@alg_router.put("/quest_person/{idQuestPerson}",tags=['Quest Person'])
def edit_quest_person(idQuestPerson:int,item:editQuestPerson,DB:Session = Depends(get_session)):
    questperson = DB.query(QuestPerson).filter(QuestPerson.id_quest_persons == idQuestPerson).first()
    if questperson is None:
        return JSONResponse(status_code= 404, content={"message":"Допрашиваемый не найден"})
    individual = DB.query(Invidivual).filter(Invidivual.individuals_id == questperson.individuals_id).first()
    passport = DB.query(Passport).filter(Passport.passports_id == individual.FK_passports_id).first()
    passport_unique = DB.query(Passport).filter(Passport.passports_id != passport.passports_id,Passport.passport_serial == item.passport_serial,Passport.passport_number == item.passport_number).first()
    if passport_unique is not None:
        return JSONResponse(status_code= 404, content={"message":"Данный паспорт уже зарегистрирован в системе"})
    individual.first_name=item.firstname
    individual.middle_name=item.middlename
    individual.last_name=item.lastname
    passport.birthdate=item.datebrithday
    passport.place_of_birth=item.bitrh_residence
    passport.place_of_residence=item.place_of_recidense
    individual.resident=item.resident
    individual.education=item.education
    individual.family_status=item.family_status
    individual.place_work=item.place_work
    individual.military=item.military
    individual.criminal=item.criminal
    passport.passport_serial=item.passport_serial
    passport.passport_number=item.passport_number
    individual.other=item.other
    questperson.notes=item.notes
    questperson.petition=item.petition
    try:
        DB.commit()
        DB.refresh(passport)
        DB.commit()
        DB.refresh(individual)
        DB.commit()
        DB.refresh(questperson)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
@alg_router.get("/quest_person/{idCase}/{idProceduralPosition}",tags=['Quest Person'])
def get_quest_person(idCase:int,idProceduralPosition:int,DB:Session = Depends(get_session)):
    if idProceduralPosition == 1 or idProceduralPosition == 8:
        quest_person = DB.query(QuestPerson).filter(QuestPerson.fk_active_case_id==idCase,QuestPerson.fk_procedural_position==idProceduralPosition).first()
        if quest_person is None:
            return JSONResponse(status_code= 404, content={"message":"Данные не найдены"})
        individual = DB.query(Invidivual).filter(Invidivual.individuals_id == quest_person.individuals_id).first()
        passport = DB.query(Passport).filter(Passport.passports_id == individual.FK_passports_id).first()
        slov = {
            "id_quest_persons":quest_person.id_quest_persons,
            "firstname":individual.first_name,
            "middlename":individual.middle_name,
            "lastname":individual.last_name,
            "datebrithday":passport.birthdate,
            "bitrh_residence":passport.place_of_birth,
            "place_of_recidense":passport.place_of_residence,
            "resident":individual.resident,
            "education":individual.education,
            "family_status":individual.family_status,
            "place_work":individual.place_work,
            "military":individual.military,
            "criminal":individual.criminal,
            "passport_serial":passport.passport_serial,
            "passport_number":passport.passport_number,
            "other":individual.other,
            "notes":quest_person.notes,
            "petition":quest_person.petition

        }
    elif idProceduralPosition == 7 or idProceduralPosition==6:
        quest_person=DB.query(QuestPerson).filter(QuestPerson.fk_active_case_id == idCase,QuestPerson.fk_procedural_position==idProceduralPosition).all()
        if len(quest_person)==0:
            return JSONResponse(status_code= 404, content={"message":"Данные не найдены"})
        slov = []
        for i in range(len(quest_person)):
            individual = DB.query(Invidivual).filter(Invidivual.individuals_id == quest_person[i].individuals_id).first()
            passport = DB.query(Passport).filter(Passport.passports_id == individual.FK_passports_id).first()
            th_slov = {
                "id_quest_persons":quest_person[i].id_quest_persons,
                "firstname":individual.first_name,
                "middlename":individual.middle_name,
                "lastname":individual.last_name,
                "datebrithday":passport.birthdate,
                "bitrh_residence":passport.place_of_birth,
                "place_of_recidense":passport.place_of_residence,
                "resident":individual.resident,
                "education":individual.education,
                "family_status":individual.family_status,
                "place_work":individual.place_work,
                "military":individual.military,
                "criminal":individual.criminal,
                "passport_serial":passport.passport_serial,
                "passport_number":passport.passport_number,
                "other":individual.other,
                "notes":quest_person[i].notes,
                "petition":quest_person[i].petition
            }
            slov.append(th_slov)
    return slov
@alg_router.get("/quest_person_individual/{idCase}",tags=['Quest Person'])
def get_quest_person(idCase:int,DB:Session = Depends(get_session)):
    statement_id = DB.query(Case).filter(Case.active_cases_id==idCase).first().FK_statement_id
    if statement_id is None:
        return JSONResponse(status_code= 404, content={"message":"Данные не найдены"})
    individual_id = DB.query(Statement).filter(Statement.statement_id == statement_id).first().fk_individuals_id
    individual = DB.query(Invidivual).filter(Invidivual.individuals_id==individual_id).first()
    passport = DB.query(Passport).filter(Passport.passports_id == individual.FK_passports_id).first()
    passport_date_issue = passport.passport_date_of_issue.strftime("%d.%m.%Y")
    person = {
        "firstname":individual.first_name,
        "middlename":individual.middle_name,
        "lastname":individual.last_name,
        "datebrithday":passport.birthdate,
        "bitrh_residence":passport.place_of_birth,
        "place_of_recidense":passport.place_of_residence,
        "passport_serial": passport.passport_serial,
        "passport_number": passport.passport_number,
        "resident":individual.resident,
        "education":individual.education,
        "family_status":individual.family_status,
        "place_work":individual.place_work,
        "military":individual.military,
        "criminal":individual.criminal,
        "other":individual.other,
    }
    return person

@alg_router.post("/person_involved_quest/",tags=['Quest'])
def create_person_involved_quest(item:createInspectPersonInvolved,DB:Session = Depends(get_session)):
    questperson = DB.query(QuestPersonInvolved).filter(QuestPersonInvolved.id_quest==item.idInspect,QuestPersonInvolved.id_persons_involved==item.idPerson).first()
    if questperson is not None:
        return JSONResponse(status_code=400, content={"message":"Данные уже имеются"})
    questperson = QuestPersonInvolved(id_persons_involved=item.idPerson,id_quest=item.idInspect,notes=item.notes,petition=item.petition,fk_procedural_position=item.fk_procedural_position)
    try:
        DB.add(questperson)
        DB.commit()
        DB.refresh(questperson)
        return JSONResponse(status_code=200, content={"message": "Данные добавлены!"})
    except HTTPException:
        return JSONResponse(status_code=500, content={"message":"Ошибка"})
@alg_router.put("/person_involved_quest/{idQuest}/{idPersonInvolved}",tags=['Quest'])
def edit_person_involved_quest_by_id(idQuest:int,idPersonInvolved:int,item:updatePersonInvolvedQuest,DB:Session = Depends(get_session)):
    quest = DB.query(Quest).filter(idQuest == Quest.id_quest).first()
    if quest is None:
        return JSONResponse(status_code=404, content={"message":"Данные о допросе не найдены"})
    person_involved = DB.query(Invidivual).filter(idPersonInvolved == Invidivual.individuals_id).first()
    if person_involved is None:
        return JSONResponse(status_code=404, content={"message":"Данные об учавствующем лице не найдены"})
    person_involved_quest = DB.query(QuestPersonInvolved).filter(quest.id_quest == QuestPersonInvolved.id_quest,person_involved.individuals_id == QuestPersonInvolved.id_persons_involved).first()
    if person_involved_quest is None:
        return JSONResponse(status_code=404, content={"message":"Данные об учавствующем лице не найдены"})
    person_involved_quest.notes = item.notes
    person_involved_quest.petition = item.petition
    try:
        DB.commit()
        DB.refresh(person_involved_quest)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
@alg_router.delete("/person_involved_quest/{idQuest}/{idPersonInvolved}",tags=['Quest'])
def delete_person_involved_quest_by_id(idQuest:int,idPersonInvolved:int,DB:Session = Depends(get_session)):
    quest = DB.query(Quest).filter(idQuest == Quest.id_quest).first()
    if quest is None:
        return JSONResponse(status_code=404, content={"message":"Данные о допросе не найдены"})
    person_involved = DB.query(Invidivual).filter(idPersonInvolved == Invidivual.individuals_id).first()
    if person_involved is None:
        return JSONResponse(status_code=404, content={"message":"Данные об учавствующем лице не найдены"})
    person_involved_quest = DB.query(QuestPersonInvolved).filter(quest.id_quest == QuestPersonInvolved.id_quest,person_involved.individuals_id == QuestPersonInvolved.id_persons_involved).first()
    if person_involved_quest is None:
        return JSONResponse(status_code=404, content={"message":"Данные об учавствующем лице не найдены"})
    try:
        DB.delete(person_involved_quest)
        DB.commit()
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
@alg_router.get("/person_involved_quest/{idQuestPerson}",tags=['Quest'])
def get_all_person_quest(idQuestPerson:int,DB:Session = Depends(get_session)):
    quest= DB.query(Quest).filter(idQuestPerson == Quest.fk_id_quest_person).first()
    if quest is None:
        return JSONResponse(status_code=404, content={"message":"Данные о допросе не найдены"})
    questperson = DB.query(QuestPersonInvolved).filter(QuestPersonInvolved.id_quest==quest.id_quest).all()
    if len(questperson) == 0:
        return JSONResponse(status_code=404,content={"message":"Данные не найдены"})
    all_person = []
    for i in range(len(questperson)):
        person = DB.query(Invidivual).filter(Invidivual.individuals_id == questperson[i].id_persons_involved).first()
        passport = DB.query(Passport).filter(Passport.passports_id == person.FK_passports_id).first()
        if questperson[i].petition == 1:
            petition = 'Нет'
        else:
            petition = 'Да'
        procedural_position_name = DB.query(ProceduralPosition).filter(questperson[i].fk_procedural_position==ProceduralPosition.id_procedural_position).first().name_position
        info = {
            "id_persons_involved":person.individuals_id,
            "first_name":person.first_name,
            "middle_name":person.middle_name,
            "last_name":person.last_name,
            "passport_serial":passport.passport_serial,
            "passport_number":passport.passport_number,
            "residence":passport.place_of_residence,
            "fk_procedural_position":procedural_position_name,
            "notes":questperson[i].notes,
            "petition":petition
        }
        all_person.append(info)
    return all_person
@alg_router.post("/quest/",tags=['Quest'])
def quest_create(item:createQuest,DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == item.user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == item.case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    if cases.fk_user_id != user.id:
        return JSONResponse(status_code= 400, content={"message": "Этот пользователь не может редактировать данное дело"})
    quest = DB.query(Quest).filter(Quest.fk_id_quest_person == item.quest_person).first()
    if quest is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    quest = Quest(date_quest=item.date_quest,begin_quest=item.begin_quest,end_quest=item.end_quest,quest_place=item.quest_place,technical_means=item.technical_means,establish=item.establish,photography=item.photography,playing=item.playing,familiarization=item.familiarization,fk_id_quest_person=item.quest_person)
    if quest is None:
        return JSONResponse(status_code= 404, content={"message": "Произошла ошибка!"})
    try:
        DB.add(quest)
        DB.commit()
        DB.refresh(quest)
        return JSONResponse(status_code= 200, content={"message": "Данные успешно добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/quest/{questPersonId}",tags=['Quest'])
def get_quest(questPersonId:int,DB:Session = Depends(get_session)):
    questperson = DB.query(QuestPerson).filter(QuestPerson.id_quest_persons == questPersonId).first()
    if questperson is None:
        return JSONResponse(status_code= 404, content={"message": "Человек не найден"})
    caseId = questperson.fk_active_case_id
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    quest = DB.query(Quest).filter(Quest.fk_id_quest_person==questPersonId).first()
    if quest is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return quest
@alg_router.put("/quest/{questPersonId}",tags=['Quest'])
def edit_quest(questPersonId:int,item:updateQuest,DB:Session = Depends(get_session)):
    questperson = DB.query(QuestPerson).filter(QuestPerson.id_quest_persons == questPersonId).first()
    if questperson is None:
        return JSONResponse(status_code= 404, content={"message": "Человек не найден"})
    caseId = questperson.fk_active_case_id
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    quest = DB.query(Quest).filter(Quest.fk_id_quest_person==questPersonId).first()
    if quest is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    quest.date_quest = item.date_quest
    quest.begin_quest = item.begin_quest
    quest.end_quest = item.end_quest
    quest.quest_place = item.quest_place
    quest.technical_means = item.technical_means
    quest.establish = item.establish
    quest.photography = item.photography
    quest.playing = item.playing
    quest.familiarization = item.familiarization
    try:
        DB.commit()
        DB.refresh(quest)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})

@alg_router.post("/set/",tags=['Set'])
def set_create(item:createSet,DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == item.user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == item.case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    if cases.fk_user_id != user.id:
        return JSONResponse(status_code= 400, content={"message": "Этот пользователь не может редактировать данное дело"})
    set = DB.query(Set).filter(Set.fk_active_case_id == item.case_id).first()
    if set is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    set = Set(fk_active_case_id=item.case_id,criminal=item.criminal)
    if set is None:
        return JSONResponse(status_code= 404, content={"message": "Произошла ошибка!"})
    try:
        DB.add(set)
        DB.commit()
        DB.refresh(set)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/set/{caseId}",tags=['Set'])
def get_set(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    set = DB.query(Set).filter(Set.fk_active_case_id==caseId).first()
    if set is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return set
@alg_router.put("/set/{caseId}",tags=['Set'])
def edit_set(caseId:int,item:updateSet,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    set = DB.query(Set).filter(Set.fk_active_case_id==caseId).first()
    if set is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    set.criminal=item.criminal
    try:
        DB.commit()
        DB.refresh(set)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})   
@alg_router.post("/declaration_victim/",tags=['Declaration Victim'])
def declaration_victim_create(item:createDeclarationVictim,DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == item.user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == item.case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    if cases.fk_user_id != user.id:
        return JSONResponse(status_code= 400, content={"message": "Этот пользователь не может редактировать данное дело"})
    victim = DB.query(DeclarationVictim).filter(DeclarationVictim.fk_active_case_id == item.case_id).first()
    if victim is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    victim = DeclarationVictim(fk_active_case_id=item.case_id,rationale=item.rationale,type_of_harm=item.type_of_harm,date_declaration=item.date_declaration,fk_individuals_id=item.fk_individuals_id)
    if victim is None:
        return JSONResponse(status_code= 404, content={"message": "Произошла ошибка!"})
    try:
        DB.add(victim)
        DB.commit()
        DB.refresh(victim)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/declaration_victim/{caseId}",tags=['Declaration Victim'])
def get_declaration_victim(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    victim = DB.query(DeclarationVictim).filter(DeclarationVictim.fk_active_case_id==caseId).first()
    if victim is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return victim
@alg_router.put("/declaration_victim/{caseId}",tags=['Declaration Victim'])
def edit_declaration_victim(caseId:int,item:updateDeclarationVictim,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    victim = DB.query(DeclarationVictim).filter(DeclarationVictim.fk_active_case_id==caseId).first()
    if victim is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    victim.rationale=item.rationale
    victim.type_of_harm=item.type_of_harm
    victim.date_declaration=item.date_declaration
    victim.fk_individuals_id = item.fk_individuals_id
    try:
        DB.commit()
        DB.refresh(victim)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
    
@alg_router.post("/declaration_criminal_case/",tags=['Declaration Criminal Case'])
def declaration_criminal_case_create(item:createDeclarationCriminalCase,DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == item.user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == item.case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    if cases.fk_user_id != user.id:
        return JSONResponse(status_code= 400, content={"message": "Этот пользователь не может редактировать данное дело"})
    criminal_case = DB.query(DeclarationCriminalCase).filter(DeclarationCriminalCase.fk_active_case_id == item.case_id).first()
    if criminal_case is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    criminal_case = DeclarationCriminalCase(fk_active_case_id=item.case_id,time_declaration=item.time_declaration,date_declaration=item.date_declaration,reason_declaration=item.reason_declaration,item_part_article=item.item_part_article,know_person=item.know_person,firstname=item.firstname,middlename=item.middlename,lastname=item.lastname,name_procurator=item.name_procurator,date_procurator=item.date_procurator,time_procurator=item.time_procurator,date_victim=item.date_victim,date_criminal_person=item.date_criminal_person,number_cases=item.number_cases)
    if criminal_case is None:
        return JSONResponse(status_code= 404, content={"message": "Произошла ошибка!"})
    try:
        DB.add(criminal_case)
        DB.commit()
        DB.refresh(criminal_case)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/declaration_criminal_case/{caseId}",tags=['Declaration Criminal Case'])
def get_declaration_criminal_case(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    criminal_case = DB.query(DeclarationCriminalCase).filter(DeclarationCriminalCase.fk_active_case_id==caseId).first()
    if criminal_case is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return criminal_case
@alg_router.put("/declaration_criminal_case/{caseId}",tags=['Declaration Criminal Case'])
def edit_declaration_victim(caseId:int,item:updateDeclarationCriminalCase,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    criminal_case = DB.query(DeclarationCriminalCase).filter(DeclarationCriminalCase.fk_active_case_id==caseId).first()
    if criminal_case is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    criminal_case.time_declaration=item.time_declaration
    criminal_case.date_declaration=item.date_declaration
    criminal_case.reason_declaration=item.reason_declaration
    criminal_case.item_part_article=item.item_part_article
    criminal_case.know_person=item.know_person
    criminal_case.firstname=item.firstname
    criminal_case.middlename=item.middlename
    criminal_case.lastname=item.lastname
    criminal_case.name_procurator=item.name_procurator
    criminal_case.date_procurator=item.date_procurator
    criminal_case.time_procurator=item.time_procurator
    criminal_case.date_victim=item.date_victim
    criminal_case.date_criminal_person=item.date_criminal_person
    criminal_case.number_cases = item.number_cases
    try:
        DB.commit()
        DB.refresh(criminal_case)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})

@alg_router.post("/order_wanted/",tags=['Order Wanted'])
def order_wanted_create(item:createOrderWanted,DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == item.user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == item.case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    if cases.fk_user_id != user.id:
        return JSONResponse(status_code= 400, content={"message": "Этот пользователь не может редактировать данное дело"})
    order_wanted = DB.query(OrderWanted).filter(OrderWanted.fk_active_case_id == item.case_id).first()
    if order_wanted is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    order_wanted = OrderWanted(fk_active_case_id=item.case_id,check_criminal=item.check_criminal,text_order_wanted=item.text_order_wanted,name_orders=item.name_orders)
    if order_wanted is None:
        return JSONResponse(status_code= 404, content={"message": "Произошла ошибка!"})
    try:
        DB.add(order_wanted)
        DB.commit()
        DB.refresh(order_wanted)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/order_wanted/{caseId}",tags=['Order Wanted'])
def get_order_wanted(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    order_wanted = DB.query(OrderWanted).filter(OrderWanted.fk_active_case_id==caseId).first()
    if order_wanted is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return order_wanted
@alg_router.put("/order_wanted/{caseId}",tags=['Order Wanted'])
def edit_order_wanted(caseId:int,item:updateOrderWanted,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    order_wanted = DB.query(OrderWanted).filter(OrderWanted.fk_active_case_id==caseId).first()
    if order_wanted is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    order_wanted.check_criminal=item.check_criminal
    order_wanted.text_order_wanted=item.text_order_wanted
    order_wanted.name_orders = item.name_orders
    try:
        DB.commit()
        DB.refresh(order_wanted)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
@alg_router.post("/request/domain/download/{id}",tags=['Request'])
def create_domain_request_document(id:int,item:requestDomain,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id==id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    flag = 0
    declaration_criminal = DB.query(DeclarationCriminalCase).filter(DeclarationCriminalCase.fk_active_case_id == cases.active_cases_id).first()
    if declaration_criminal is None:
        flag = 1
    individuals_id=DB.query(Statement).filter(Statement.statement_id==cases.FK_statement_id).first().fk_individuals_id
    individual = DB.query(Invidivual).filter(Invidivual.individuals_id==individuals_id).first()
    statement = DB.query(Statement).filter(Statement.statement_id == cases.FK_statement_id).first()
    statement_date = statement.statement_date.strftime("%d.%m.%Yг.")
    morph = pymorphy3.MorphAnalyzer()
    passport = DB.query(Passport).filter(Passport.passports_id == individual.FK_passports_id).first()
    date_birthday = passport.birthdate.strftime("%d.%m.%Y")
    individualFIO_2 = individual.first_name[0] + morph.parse(individual.first_name)[0].inflect({'accs'}).word[1:] + ' ' + individual.middle_name[0]+'.' + individual.last_name[0]+'.'
    if morph.parse(item.middle_name)[0].tag.gender == 'femn':
        obr = 'Уважаемая'
    else:
        obr = "Уважаемый"
    if flag == 0:
        date_case = declaration_criminal.date_declaration.strftime("%d.%m.%Y г.")
        number_cases = declaration_criminal.number_cases
        docname = 'requestdomain.docx'
    else:
        date_case = "sad"
        number_cases = "asd"
        docname = 'requestdomain_2.docx'
    dst = f'media/{cases.fk_user_id}/cases/{cases.active_cases_id}/documents/'    
    src = 'documents/' + docname
    dstsrc = 'Запрос оператору доменного имени.docx'
    shutil.copy(src, dst, follow_symlinks=True)
    if os.path.isfile(dst+'Запрос оператору доменного имени.docx'):
        os.remove(dst+'Запрос оператору доменного имени.docx')
    os.rename(dst+docname,dst+'Запрос оператору доменного имени.docx')
    doc = docx.Document(dst+dstsrc)
    style1 = doc.styles['Normal']
    style1.font.name = 'Times New Roman'
    style1.font.size = Pt(13)
    dictionary = {
        "no":item.org_name,
        "nr": item.middle_name[0]+'.'+item.last_name[0]+'. '+item.first_name[0]+morph.parse(item.first_name)[0].inflect({'datv'}).word[1:],
        "ad": item.adress,
        "КУСП": 'КУСП' + ' № ' + statement.number_statement + ' от ' + statement_date,
        "FIOZ": individualFIO_2 + ', '+date_birthday +' года рождения, серия паспорта - '+passport.passport_serial+', номер паспорта - '+ passport.passport_number,
        "****":item.money,
        "numbcase":number_cases,
        "dcase": date_case,
        "dn": item.domain_name,
        "obr":obr,
        "io": item.middle_name + ' ' + item.last_name +'!'
    }
    for i in dictionary:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i,dictionary[i])
    doc.save(dst+dstsrc)
    return FileResponse(dst+dstsrc,filename="Запрос оператору доменного имени.docx",media_type='application/octet-stream')
@alg_router.post("/request/mobile_operator/download/{id}",tags=['Request'])
def create_mobile_request_document(id:int,item:requestMobile,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id==id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    flag = 0
    declaration_case = DB.query(DeclarationCriminalCase).filter(DeclarationCriminalCase.fk_active_case_id == cases.active_cases_id).first()
    if declaration_case is None:
        flag = 1
    individuals_id=DB.query(Statement).filter(Statement.statement_id==cases.FK_statement_id).first().fk_individuals_id
    individual = DB.query(Invidivual).filter(Invidivual.individuals_id==individuals_id).first()
    statement = DB.query(Statement).filter(Statement.statement_id == cases.FK_statement_id).first()
    statement_date = statement.statement_date.strftime("%d.%m.%Yг.")
    morph = pymorphy3.MorphAnalyzer()
    passport = DB.query(Passport).filter(Passport.passports_id == individual.FK_passports_id).first()
    date_birthday = passport.birthdate.strftime("%d.%m.%Y")
    individualFIO = individual.first_name[0] + morph.parse(individual.first_name)[0].inflect({'datv'}).word[1:] + ' ' + individual.middle_name[0]+'.' + individual.last_name[0]+'.'
    individualFIO_2 = individual.first_name[0] + morph.parse(individual.first_name)[0].inflect({'accs'}).word[1:] + ' ' + individual.middle_name[0]+'.' + individual.last_name[0]+'.'
    user = DB.query(User).filter(User.id==cases.fk_user_id).first()
    userFIO = user.middle_name[0]+'.' + user.last_name[0]+'.'+ ' ' + user.first_name
    individual_sex = DB.query(Passport).filter(Passport.passports_id==individual.FK_passports_id).first().sex
    department = DB.query(Department).filter(user.fk_department_id == Department.departments_id).first()
    name_department = department.name_department
    victim = 'потерпевшей'
    if individual_sex=='Муж':
        victim = 'потерпевшему'
    if flag == 0:
        date_case = declaration_case.date_declaration.strftime("%d.%m.%Y")
        number_cases = declaration_case.number_cases
        docname = 'requestmobile.docx'
    else:
        date_case = "sad"
        number_cases="sad"
        docname = 'requestmobile_2.docx'
    date_incidient = item.date_incidient.strftime("%d.%m.%Y")
    src = 'documents/' + docname
    dst = f'media/{cases.fk_user_id}/cases/{cases.active_cases_id}/documents/'
    shutil.copy(src, dst, follow_symlinks=True)
    if os.path.isfile(dst+'Запрос мобильному оператору.docx'):
        os.remove(dst+'Запрос мобильному оператору.docx')
    os.rename(dst+docname,dst+'Запрос мобильному оператору.docx')
    doc = docx.Document(dst+'Запрос мобильному оператору.docx')
    dictionary = {"nop":item.name_org,
                  "СОВЕТСКИЙ":name_department.upper(),"Советский":department.name_department,
                  "КУСП": 'КУСП' + ' № ' + statement.number_statement + ' от ' + statement_date,
                  "FIOZ": individualFIO_2 + ', '+date_birthday +' года рождения, серия паспорта - '+passport.passport_serial+', номер паспорта - '+ passport.passport_number,
                  "****":item.money,
                  "Челябинск":department.city,"ЧЕЛЯБИНСК":department.city.upper(),
                  "ул. Монакова 2":department.street+', '+department.house_number,"454091":department.postal_code,"1220175009800000":number_cases,
                  "30.01.2022":date_case,"Сидоровой А.В.":individualFIO,
                  "13 часов":item.time[0]+item.time[1]+' часов',"31 минуту":item.time[3]+item.time[4]+' минуту',
                  "21 января 2022":date_incidient,"8-499-769-34-13":item.phone,"11 минут 24 секунды":item.duration,
                  "ivanova@mvd.ru":user.email,"И.И. Иванова":userFIO,"потерпевшей":victim,"Старший следователь":user.appointment,
                  "старший лейтенант":user.users_rank}
    for i in dictionary:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i,dictionary[i])
    doc.save(dst+'Запрос мобильному оператору.docx')
    return FileResponse(dst+'Запрос мобильному оператору.docx',filename="Запрос мобильному оператору.docx",media_type='application/octet-stream')
@alg_router.post("/request/bank/download/{id}",tags=['Request'])
def create_bank_request_document(id:int,item:requestBank,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id==id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    flag = 0
    declaration_cases = DB.query(DeclarationCriminalCase).filter(DeclarationCriminalCase.fk_active_case_id == cases.active_cases_id).first()
    if declaration_cases is None:
        flag = 1
        item_part_article = "sad"
    else: 
        item_part_article = declaration_cases.item_part_article
    individuals_id=DB.query(Statement).filter(Statement.statement_id==cases.FK_statement_id).first().fk_individuals_id
    individual = DB.query(Invidivual).filter(Invidivual.individuals_id==individuals_id).first()
    statement = DB.query(Statement).filter(Statement.statement_id == cases.FK_statement_id).first()
    statement_date = statement.statement_date.strftime("%d.%m.%Yг.")
    morph = pymorphy3.MorphAnalyzer()
    passport = DB.query(Passport).filter(Passport.passports_id == individual.FK_passports_id).first()
    date_birthday = passport.birthdate.strftime("%d.%m.%Y")
    individualFIO = individual.first_name[0] + morph.parse(individual.first_name)[0].inflect({'accs'}).word[1:] + ' ' + individual.middle_name[0]+'.' + individual.last_name[0]+'.'
    user = DB.query(User).filter(User.id==cases.fk_user_id).first()
    userFIO = user.middle_name[0]+'.' + user.last_name[0]+'.'+ ' ' + user.first_name
    userFIO2 = user.middle_name[0]+'.' + user.last_name[0]+'.'+ ' ' + user.first_name[0] + morph.parse(user.first_name)[0].inflect({'gent'}).word[1:]
    mainUser = DB.query(User).filter(User.FK_user_roles_id==4,User.fk_department_id==user.fk_department_id).first()
    mainUserFIO = mainUser.middle_name[0]+'.' + mainUser.last_name[0]+'.'+ ' ' + mainUser.first_name
    department = DB.query(Department).filter(user.fk_department_id == Department.departments_id).first()
    name_department = department.name_department
    src = 'documents/'
    if flag == 0:
        docname = 'requestbank.docx'
        number_cases = declaration_cases.number_cases
    else:
        docname = 'requestbank_2.docx'
        number_cases = "asd"
    src = src + docname
    dst = f'media/{cases.fk_user_id}/cases/{cases.active_cases_id}/documents/'
    shutil.copy(src, dst, follow_symlinks=True)
    if os.path.isfile(dst+'Запрос в банк.docx'):
        os.remove(dst+'Запрос в банк.docx')
    os.rename(dst+docname,dst+'Запрос в банк.docx')
    doc = docx.Document(dst+'Запрос в банк.docx')
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    dictionary = {"СОВЕТСКИЙ":name_department.upper(), "Советский":department.name_department,
                  "Челябинск":department.city,"ЧЕЛЯБИНСК":department.city.upper(),
                  "ул. Монакова 2":department.street+', '+department.house_number,"454091":department.postal_code,
                  "КУСП": 'КУСП' + ' № ' + statement.number_statement + ' от ' + statement_date,
                  "FIOZ": individualFIO + ', '+date_birthday +' года рождения, серия паспорта - '+passport.passport_serial+', номер паспорта - '+ passport.passport_number,
                  "****":item.money,
                  "1220175009800000":number_cases,
                  "АО «Альфа Банк»":item.bank,
                  "pr":item_part_article,
                  "Сидоровой А.В.":individualFIO,"ivanovamvd@.ru":user.email,"И.И. Иванова":userFIO,"Старший следователь":user.appointment,"старший лейтенант":user.users_rank,
                  "К.А. Борисова":mainUserFIO, "подполковник":mainUser.users_rank,"И.И.  Ивановой":userFIO2}
    
    for i in dictionary:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i,dictionary[i])
    doc.save(dst+'Запрос в банк.docx')
    return FileResponse(dst+'Запрос в банк.docx',filename="Запрос в банк.docx",media_type='application/octet-stream')
@alg_router.get("/documents/inspect/{id}",tags=['Inspect'])
def get_inspect_document(id:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id==id).first()
    if cases is None:
        return JSONResponse(status_code=404,content={"message":"Дело не найдено"})
    department_id = DB.query(User).filter(User.id==cases.fk_user_id).first().fk_department_id
    department = DB.query(Department).filter(Department.departments_id==department_id).first()
    inspect = DB.query(Inspect).filter(Inspect.fk_active_case_id == cases.active_cases_id).first()
    date_inspect = inspect.date_inspect.strftime("%d.%m.%Y")
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    userFIO_1 = user.first_name + ' ' + user.middle_name[0]+'.' + user.last_name[0]+'.'
    userFIO_2 = user.middle_name[0]+'.' + user.last_name[0]+'.'+ ' ' + user.first_name
    inspect_people = DB.query(InspectPersonsInvolved).filter(inspect.id_inspect == InspectPersonsInvolved.id_inspect).all()
    p1 = ''
    p2 = ''
    pet = ''
    notes = ''
    otherPerson = ''
    morph = pymorphy3.MorphAnalyzer()
    for i in range(len(inspect_people)):
        person = DB.query(Invidivual).filter(Invidivual.individuals_id == inspect_people[i].id_persons_involved).first()
        passport = DB.query(Passport).filter(Passport.passports_id == person.FK_passports_id).first()
        procedural_position_name = DB.query(ProceduralPosition).filter(inspect_people[i].fk_procedural_position==ProceduralPosition.id_procedural_position).first().name_position
        if procedural_position_name == 'Понятой':
            if len(p1) == 0:
                p1 = person.first_name+' '+person.middle_name+' '+person.last_name+'. Место жительства: '+ passport.place_of_residence
            else:
                p2 = person.first_name+' '+person.middle_name+' '+person.last_name+'. Место жительства: '+ passport.place_of_residence
        elif procedural_position_name == 'Эксперт':
            expFIO = person.first_name[0]+morph.parse(person.first_name)[0].inflect({'datv','sing'}).word[1:] + ' ' + person.middle_name[0]+morph.parse(person.middle_name)[0].inflect({'datv','sing'}).word[1:] + ' ' + person.last_name[0]+morph.parse(person.last_name)[0].inflect({'datv','sing'}).word[1:]
            exp = 'Эксперт'
            state = '57'
        elif procedural_position_name == 'Специалист':
            expFIO = person.first_name[0]+morph.parse(person.first_name)[0].inflect({'datv','sing'}).word[1:] + ' ' + person.middle_name[0]+morph.parse(person.middle_name)[0].inflect({'datv','sing'}).word[1:] + ' ' + person.last_name[0]+morph.parse(person.last_name)[0].inflect({'datv','sing'}).word[1:]
            exp = 'Специалист'
            state = '58'
        else:
            if len(otherPerson) > 0:
                otherPerson += ', '
            otherPerson += morph.parse(procedural_position_name.lower())[0].inflect({'gent'}).word + ' '+ person.first_name[0]+morph.parse(person.first_name)[0].inflect({'gent'}).word[1:] + ' ' + person.middle_name[0]+morph.parse(person.middle_name)[0].inflect({'gent'}).word[1:] + ' ' + person.last_name[0]+morph.parse(person.last_name)[0].inflect({'gent'}).word[1:]
        if inspect_people[i].petition == 2:
            if len(pet) > 0:
                pet += ', '
            pet += morph.parse(procedural_position_name.lower())[0].inflect({'gent','masc'}).word +' ' + person.first_name[0] + morph.parse(person.first_name)[0].inflect({'gent'}).word[1:] + ' '+person.middle_name[0]+'.'+person.last_name[0]+'.'
        if len(inspect_people[i].notes) != 0:
            if len(notes) > 0:
                notes += ', '
            notes += procedural_position_name.lower()+' ' + person.first_name + ' '+person.middle_name[0]+'.'+person.last_name[0]+'.' +' Замечание - ' +inspect_people[i].notes+'.'
    if pet == '':
        pet = 'Следователя'+userFIO_1
    if len(notes) != 0:
        notes ='сделали - ' + notes
    else:
        notes = 'не сделали.'
    src = 'documents/omp.docx'
    dst = f'media/{cases.fk_user_id}/cases/{cases.active_cases_id}/documents/'
    shutil.copy(src, dst, follow_symlinks=True)
    if os.path.isfile(dst+'Протокол осмотра места происшествия.docx'):
        os.remove(dst+'Протокол осмотра места происшествия.docx')
    os.rename(dst+'omp.docx',dst+'Протокол осмотра места происшествия.docx')
    doc = docx.Document(dst+'Протокол осмотра места происшествия.docx')
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    dictionary = {
        "othp": otherPerson,
        "inex": inspect.inspect_exam.lower(),
        "pet": pet,
        "notes": notes,
        "p1":p1,
        "p2":p2,
        "Exs":exp,
        "exFIO":expFIO,
        "st": state,
        "App": user.appointment,
        "ur": user.users_rank.lower(),
        "FIO1": userFIO_1,
        "FIO2": userFIO_2,
        "(мс)": "г. "+department.city,
        "Домп": date_inspect,
        "омпнч": inspect.begin_inspect[0]+inspect.begin_inspect[1],
        "онм": inspect.begin_inspect[3]+inspect.begin_inspect[4],
        "ооч": inspect.end_inspect[0]+inspect.end_inspect[1],
        "оом": inspect.end_inspect[3]+inspect.end_inspect[4],
        "окс": inspect.from_message_inspect.lower(),
        "очс": inspect.message_inspect.lower(),
        "кпс": inspect.place_of_inspect.lower(),
        "tm": inspect.technical_means.lower(),
        "уос": inspect.conditions.lower(),
        "оуе": inspect.establish.lower(),
        "фва": inspect.photography.lower(),
        "si": inspect.sized_items.lower(),
        "ifi": inspect.items_for_inspect.lower(),
        "fam":inspect.familiarization.lower()
    }
    all_tables = doc.tables
    
    for i in dictionary:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i,dictionary[i])
        for k, table in enumerate(all_tables):
            for j, row in enumerate(table.rows):
                for cell in row.cells:
                    if cell.text.find(i) >= 0:
                        cell.text = cell.text.replace(i,dictionary[i])
    doc.save(dst+'Протокол осмотра места происшествия.docx')
    return FileResponse(dst+'Протокол осмотра места происшествия.docx',filename="Протокол осмотра места происшествия.docx",media_type='application/octet-stream')
@alg_router.get("/documents/check/{id}",tags=['Check'])
def get_check_document(id:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id==id).first()
    if cases is None:
        return JSONResponse(status_code=404,content={"message":"Дело не найдено"})
    department_id = DB.query(User).filter(User.id==cases.fk_user_id).first().fk_department_id
    department = DB.query(Department).filter(Department.departments_id==department_id).first()
    check = DB.query(Check).filter(Check.fk_active_case_id == cases.active_cases_id).first()
    date_check = check.date_check.strftime("%d.%m.%Y")
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    userFIO_1 = user.first_name + ' ' + user.middle_name[0]+'.' + user.last_name[0]+'.'
    userFIO_2 = user.middle_name[0]+'.' + user.last_name[0]+'.'+ ' ' + user.first_name
    check_people = DB.query(CheckPersonsInvolved).filter(check.id_check == CheckPersonsInvolved.id_check).all()
    p1 = ''
    p2 = ''
    pet = ''
    notes = ''
    otherPerson = ''
    morph = pymorphy3.MorphAnalyzer()
    for i in range(len(check_people)):
        person = DB.query(Invidivual).filter(Invidivual.individuals_id == check_people[i].id_persons_involved).first()
        passport = DB.query(Passport).filter(Passport.passports_id == Invidivual.FK_passports_id).first()
        procedural_position_name = DB.query(ProceduralPosition).filter(check_people[i].fk_procedural_position==ProceduralPosition.id_procedural_position).first().name_position
        if procedural_position_name == 'Понятой':
            if len(p1) == 0:
                p1 = person.first_name+' '+person.middle_name+' '+person.last_name+'. Место жительства: '+ passport.place_of_residence
            else:
                p2 = person.first_name+' '+person.middle_name+' '+person.last_name+'. Место жительства: '+ passport.place_of_residence
        elif procedural_position_name == 'Эксперт':
            expFIO = person.first_name[0]+morph.parse(person.first_name)[0].inflect({'datv','sing'}).word[1:] + ' ' + person.middle_name[0]+morph.parse(person.middle_name)[0].inflect({'datv','sing'}).word[1:] + ' ' + person.last_name[0]+morph.parse(person.last_name)[0].inflect({'datv','sing'}).word[1:]
            exp = 'Эксперт'
            state = '57'
        elif procedural_position_name == 'Специалист':
            expFIO = person.first_name[0]+morph.parse(person.first_name)[0].inflect({'datv','sing'}).word[1:] + ' ' + person.middle_name[0]+morph.parse(person.middle_name)[0].inflect({'datv','sing'}).word[1:] + ' ' + person.last_name[0]+morph.parse(person.last_name)[0].inflect({'datv','sing'}).word[1:]
            exp = 'Специалист'
            state = '58'
        else:
            if len(otherPerson) > 0:
                otherPerson += ', '
            otherPerson += morph.parse(procedural_position_name.lower())[0].inflect({'gent'}).word + ' '+ person.first_name[0]+morph.parse(person.first_name)[0].inflect({'gent'}).word[1:] + ' ' + person.middle_name[0]+morph.parse(person.middle_name)[0].inflect({'gent'}).word[1:] + ' ' + person.last_name[0]+morph.parse(person.last_name)[0].inflect({'gent'}).word[1:]
        if check_people[i].petition == 2:
            if len(pet) > 0:
                pet += ', '
            pet += morph.parse(procedural_position_name.lower())[0].inflect({'gent','masc'}).word +' ' + person.first_name[0] + morph.parse(person.first_name)[0].inflect({'gent'}).word[1:] + ' '+person.middle_name[0]+'.'+person.last_name[0]+'.'
        if len(check_people[i].notes) != 0:
            if len(notes) > 0:
                notes += ', '
            notes += procedural_position_name.lower()+' ' + person.first_name + ' '+person.middle_name[0]+'.'+person.last_name[0]+'.' +' Замечание - ' +check_people[i].notes+'.'
    if pet == '':
        pet = 'Следователя'+userFIO_1
    if len(notes) != 0:
        notes ='сделали - ' + notes
    else:
        notes = 'не сделали.'
    src = 'documents/op.docx'
    dst = f'media/{cases.fk_user_id}/cases/{cases.active_cases_id}/documents/'
    shutil.copy(src, dst, follow_symlinks=True)
    if os.path.isfile(dst+'Протокол осмотра предметов.docx'):
        os.remove(dst+'Протокол осмотра предметов.docx')
    os.rename(dst+'op.docx',dst+'Протокол осмотра предметов.docx')
    doc = docx.Document(dst+'Протокол осмотра предметов.docx')
    style = doc.styles['ConsNonformat']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    dictionary = {
        "pd":check.items.lower(),
        "othp": otherPerson,
        "pet": pet,
        "notes": notes,
        "p1":p1,
        "p2":p2,
        "Exs":exp,
        "exFIO":expFIO,
        "st": state,
        "App": user.appointment,
        "ur": user.users_rank.lower(),
        "FIO1": userFIO_1,
        "FIO2": userFIO_2,
        "(мс)": "г. "+department.city,
        "Домп": date_check,
        "омпнч": check.begin_check[0]+check.begin_check[1],
        "онм": check.begin_check[3]+check.begin_check[4],
        "ооч": check.end_check[0]+check.end_check[1],
        "оом": check.end_check[3]+check.end_check[4],
        "whc": check.check_place.lower(),
        "watc": check.check_what.lower(),
        "wsi": check.when_size.lower(),
        "whsi": check.where_size.lower(),
        "asi": check.which_action_size.lower(),
        "(wrap)":check.wrap.lower(),
        "tm": check.technical_means.lower(),
        "уос": check.conditions.lower(),
        "оуе": check.establish.lower(),
        "фва": check.photography.lower(),
        "ifi": check.items_for_check.lower(),
        "fam":check.familiarizathion.lower()
    }
    all_tables = doc.tables
    for i in dictionary:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i,dictionary[i])
        for k, table in enumerate(all_tables):
            for j, row in enumerate(table.rows):
                for cell in row.cells:
                    if cell.text.find(i) >= 0:
                        cell.text = cell.text.replace(i,dictionary[i])
    doc.save(dst+'Протокол осмотра предметов.docx')
    return FileResponse(dst+'Протокол осмотра предметов.docx',filename="Протокол осмотра предметов.docx",media_type='application/octet-stream')

@alg_router.get("/documents/quest/{id}",tags=['Quest'])
def get_quest_victim_document(id:int,DB:Session = Depends(get_session)):
    questperson = DB.query(QuestPerson).filter(QuestPerson.id_quest_persons == id).first()
    if questperson is None:
        return JSONResponse(status_code=404,content={"message":"Допрашиваемый человек не найден"})
    questperson_individual = DB.query(Invidivual).filter(questperson.individuals_id == Invidivual.individuals_id).first()
    passport = DB.query(Passport).filter(Passport.passports_id == questperson_individual.FK_passports_id).first()
    cases = DB.query(Case).filter(Case.active_cases_id==questperson.fk_active_case_id).first()
    if cases is None:
        return JSONResponse(status_code=404,content={"message":"Дело не найдено"})
    criminal = ""
    item_part_article=""
    declarartion_criminal = DB.query(DeclarationCriminalCase).filter(DeclarationCriminalCase.fk_active_case_id == cases.active_cases_id).first()
    if declarartion_criminal is None:
        return JSONResponse(status_code=404, content={"message":"Постановление о возбуждении уголовного дела не найдено"})
    statement = DB.query(Statement).filter(cases.FK_statement_id == Statement.statement_id).first()
    criminal = statement.explanation_text
    item_part_article = declarartion_criminal.item_part_article
    department_id = DB.query(User).filter(User.id==cases.fk_user_id).first().fk_department_id
    department = DB.query(Department).filter(Department.departments_id==department_id).first()
    quest = DB.query(Quest).filter(Quest.fk_id_quest_person == questperson.id_quest_persons).first()
    if quest is None:
        return JSONResponse(status_code=404,content={"message":"Протокол допроса не найден"})
    date_quest = quest.date_quest.strftime("%d.%m.%Y")
    birthday = passport.birthdate.strftime("%d.%m.%Y")
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    userFIO_1 = user.first_name + ' ' + user.middle_name[0]+'.' + user.last_name[0]+'.'
    userFIO_2 = user.middle_name[0]+'.' + user.last_name[0]+'.'+ ' ' + user.first_name
    quest_people = DB.query(QuestPersonInvolved).filter(quest.id_quest == QuestPersonInvolved.id_quest).all()
    morph = pymorphy3.MorphAnalyzer()
    p1 = ''
    p2 = ''
    pet = ''
    notes = ''
    otherPerson = ''
    all_people_dat = questperson_individual.first_name[0] + morph.parse(questperson_individual.first_name.lower())[0].inflect({'datv'}).word[1:] +' ' + questperson_individual.middle_name[0] + '.' + questperson_individual.last_name[0] + '.'
    all_people = questperson_individual.first_name + ' ' + questperson_individual.middle_name[0] + '.' + questperson_individual.last_name[0]+ '.'
    for i in range(len(quest_people)):
        person = DB.query(Invidivual).filter(Invidivual.individuals_id == quest_people[i].id_persons_involved).first()
        procedural_position_name = DB.query(ProceduralPosition).filter(quest_people[i].fk_procedural_position==ProceduralPosition.id_procedural_position).first().name_position
        if procedural_position_name != 'Потерпевший' and procedural_position_name !='Свидетель' and procedural_position_name !='Подозреваемый':
            if len(otherPerson) == 0:
                otherPerson = ', с участием '
            elif len(otherPerson) > 0:
                otherPerson += ', '
            all_people_dat += ', ' + person.first_name[0]+morph.parse(person.first_name)[0].inflect({'datv'}).word[1:] + ' ' + person.middle_name[0] + '.' + person.last_name[0] + '.'
            all_people += ', ' + person.first_name + ' ' + person.middle_name[0] +'.' + person.last_name[0] + '.'
            otherPerson += morph.parse(procedural_position_name.lower())[0].inflect({'gent'}).word + ' '+ person.first_name[0]+morph.parse(person.first_name)[0].inflect({'gent'}).word[1:] + ' ' + person.middle_name[0]+morph.parse(person.middle_name)[0].inflect({'gent'}).word[1:] + ' ' + person.last_name[0]+morph.parse(person.last_name)[0].inflect({'gent'}).word[1:]
        if quest_people[i].petition == 2:
            if len(pet) > 0:
                pet += ', '
            pet += morph.parse(procedural_position_name.lower())[0].inflect({'gent','masc'}).word +' ' + person.first_name[0] + morph.parse(person.first_name)[0].inflect({'gent'}).word[1:] + ' '+person.middle_name[0]+'.'+person.last_name[0]+'.'
        if len(quest_people[i].notes) != 0:
            if len(notes) > 0:
                notes += ', '
            notes += procedural_position_name.lower()+' ' + person.first_name + ' '+person.middle_name[0]+'.'+person.last_name[0]+'.' +' Замечание - ' +quest_people[i].notes
    questperson_procedural_position_name = DB.query(ProceduralPosition).filter(questperson.fk_procedural_position == ProceduralPosition.id_procedural_position).first().name_position
    if questperson.petition == 2:
        if len(pet) > 0:
            pet += ', '
        pet += morph.parse(questperson_procedural_position_name.lower())[0].inflect({'gent'}).word +' ' + questperson_individual.first_name[0] + morph.parse(questperson_individual.first_name)[0].inflect({'gent'}).word[1:] + ' '+questperson_individual.middle_name[0]+'.'+questperson_individual.last_name[0]+'.'
    if pet == '':
        pet = 'следователя'+' '+user.first_name[0]+morph.parse(user.first_name)[0].inflect({'gent'}).word[1:]+ ' ' + user.middle_name[0]+'.' + user.last_name[0]+'.'
    if len(questperson.notes) !=0:
        if len(notes) > 0:
            notes += ', '
        notes += questperson_procedural_position_name.lower() +' ' + questperson_individual.first_name + ' '+questperson_individual.middle_name[0]+'.'+questperson_individual.last_name[0]+'.' +' Замечание - ' +questperson.notes
    if len(notes) != 0:
        if len(quest_people)> 0:
            notes ='сделали - ' + notes
        else:
            notes = 'сделал - ' + notes
    else:
        if len(quest_people)>0:
            notes = 'не сделали'
        else:
            notes = 'не сделал'
    if len(quest_people) > 0:
        mp = "указанным лицам"
        him = "их"
        podp = "подписями этих лиц"
    else:
        mp = "указанному лицу"
        him = "его"
        podp = "его подписью"
    dst = f'media/{cases.fk_user_id}/cases/{cases.active_cases_id}/documents/'    
    if questperson.fk_procedural_position == 1:
        src = 'documents/pdp.docx'
        dstsrc = 'pdp.docx'
    elif questperson.fk_procedural_position == 7:
        src = 'documents/pds.docx'
        dstsrc = 'pds.docx'
    elif questperson.fk_procedural_position == 6:
        src = 'documents/pdpod.docx'
        dstsrc = 'pdpod.docx'
    shutil.copy(src, dst, follow_symlinks=True)
    if questperson.fk_procedural_position == 7:
        if os.path.isfile(dst+'Допрос свидетеля '+questperson_individual.first_name+' '+questperson_individual.middle_name[0] + questperson_individual.last_name[0] +'.docx'):
            os.remove(dst+'Допрос свидетеля '+questperson_individual.first_name+' '+questperson_individual.middle_name[0] + questperson_individual.last_name[0] +'.docx')
        os.rename(dst+dstsrc,dst+'Допрос свидетеля '+questperson_individual.first_name+' '+questperson_individual.middle_name[0] + questperson_individual.last_name[0] +'.docx')
        dstsrc = 'Допрос свидетеля '+questperson_individual.first_name+' '+questperson_individual.middle_name[0] + questperson_individual.last_name[0] +'.docx'
    elif questperson.fk_procedural_position == 6:
        if os.path.isfile(dst+'Допрос подозреваемого '+questperson_individual.first_name+' '+questperson_individual.middle_name[0] + questperson_individual.last_name[0] +'.docx'):
           os.remove(dst+'Допрос подозреваемого '+questperson_individual.first_name+' '+questperson_individual.middle_name[0] + questperson_individual.last_name[0] +'.docx')
        os.rename(dst+dstsrc,dst+'Допрос подозреваемого '+questperson_individual.first_name+' '+questperson_individual.middle_name[0] + questperson_individual.last_name[0] +'.docx')
        dstsrc = 'Допрос подозреваемого '+questperson_individual.first_name+' '+questperson_individual.middle_name[0] + questperson_individual.last_name[0] +'.docx'
    else:
        if os.path.isfile(dst+'Допрос потерпевшего' +'.docx'):
           os.remove(dst+'Допрос потерпевшего' +'.docx')
        os.rename(dst+dstsrc,dst+'Допрос потерпевшего'+'.docx')
        dstsrc = 'Допрос потерпевшего' +'.docx'
    doc = docx.Document(dst+dstsrc)
    style1 = doc.styles['Normal']
    style1.font.name = 'Times New Roman'
    style1.font.size = Pt(13)
    style = doc.styles['ConsNonformat']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    dictionary = {
        "fiovictim": questperson_individual.first_name +' ' + questperson_individual.middle_name+' ' + questperson_individual.last_name,
        "birthvictim": birthday,
        "allp": all_people_dat,
        "fp": questperson_individual.first_name[0] + morph.parse(questperson_individual.first_name.lower())[0].inflect({'datv'}).word[1:] +' ' + questperson_individual.middle_name[0] + '.' + questperson_individual.last_name[0] + '.',
        "prest": criminal,
        "ukrf": item_part_article,
        "mp": mp,
        "him": him,
        "podp": podp,
        "pepli":all_people,
        "placebirth": passport.place_of_birth,
        "resvictim": passport.place_of_residence,
        "publvic": questperson_individual.resident,
        "educvic": questperson_individual.education,
        "falmvic": questperson_individual.family_status,
        "workvic": questperson_individual.place_work,
        "warvic": questperson_individual.military,
        "jailvic": questperson_individual.criminal,
        "passpvic": passport.passport_serial + ' ' + passport.passport_number,
        "othervic": questperson_individual.other,
        "othp": otherPerson,
        "pet": pet,
        "notes": notes,
        "p1":p1,
        "p2":p2,
        "App": user.appointment,
        "ur": user.users_rank.lower(),
        "FIO1": userFIO_1,
        "FIO2": userFIO_2,
        "(мс)": "г. "+department.city,
        "Домп": date_quest,
        "омпнч": quest.begin_quest[0]+quest.begin_quest[1],
        "онм": quest.begin_quest[3]+quest.begin_quest[4],
        "ооч": quest.end_quest[0]+quest.end_quest[1],
        "оом": quest.end_quest[3]+quest.end_quest[4],
        "whc": quest.quest_place.lower(),
        "numbcase": declarartion_criminal.number_cases,
        "tm": quest.technical_means.lower(),
        "оуе": quest.establish.lower(),
        "фва": quest.photography.lower(),
        "play": quest.playing.lower(),
        "fam":quest.familiarization.lower()
    }
    all_tables = doc.tables
    
    for i in dictionary:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i,dictionary[i])
        for k, table in enumerate(all_tables):
            for j, row in enumerate(table.rows):
                for cell in row.cells:
                    if cell.text.find(i) >= 0:
                        cell.text = cell.text.replace(i,dictionary[i])
    doc.save(dst+dstsrc)
    if questperson.fk_procedural_position == 1:
        return FileResponse(dst+dstsrc,filename="Протокол допроса потерпевшего.docx",media_type='application/octet-stream')
    elif questperson.fk_procedural_position == 7:
        return FileResponse(dst+dstsrc,filename="Протокол допроса свидетеля.docx",media_type='application/octet-stream')
    elif questperson.fk_procedural_position == 6:
        return FileResponse(dst+dstsrc,filename="Протокол допроса подозреваемого.docx",media_type='application/octet-stream')

@alg_router.get("/documents/declaration_victim/{id}",tags=['Declaration Victim'])
def get_declaration_victim_document(id:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id==id).first()
    if cases is None:
        return JSONResponse(status_code=404,content={"message":"Дело не найдено"})
    department_id = DB.query(User).filter(User.id==cases.fk_user_id).first().fk_department_id
    department = DB.query(Department).filter(Department.departments_id==department_id).first()
    declarartion_criminal = DB.query(DeclarationCriminalCase).filter(DeclarationCriminalCase.fk_active_case_id == cases.active_cases_id).first()
    if declarartion_criminal is None:
        return JSONResponse(status_code=404, content={"message":"Постановление о возбуждении уголовного дела не найдено"})
    declration = DB.query(DeclarationVictim).filter(DeclarationVictim.fk_active_case_id == cases.active_cases_id).first()
    if declration is None:
        return JSONResponse(status_code=404,content={"message":"Постановление не найдено"})
    individual_id = DB.query(DeclarationVictim).filter(DeclarationVictim.fk_active_case_id==cases.active_cases_id).first().fk_individuals_id
    invidivual = DB.query(Invidivual).filter(Invidivual.individuals_id == individual_id).first()
    date_declaration = declration.date_declaration.strftime("%d.%m.%Y")
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    userFIO_1 = user.first_name + ' ' + user.middle_name[0]+'.' + user.last_name[0]+'.'
    userFIO_2 = user.middle_name[0]+'.' + user.last_name[0]+'.'+ ' ' + user.first_name
    morph = pymorphy3.MorphAnalyzer()
    individual_dat = invidivual.first_name[0]+morph.parse(invidivual.first_name)[0].inflect({'datv'}).word[1:]+' '+ invidivual.middle_name[0]+morph.parse(invidivual.middle_name)[0].inflect({'datv'}).word[1:]+' '+ invidivual.last_name[0]+morph.parse(invidivual.last_name)[0].inflect({'datv'}).word[1:]
    individual_rod = invidivual.first_name[0]+morph.parse(invidivual.first_name)[0].inflect({'accs'}).word[1:]+' '+ invidivual.middle_name[0]+morph.parse(invidivual.middle_name)[0].inflect({'accs'}).word[1:]+' '+ invidivual.last_name[0]+morph.parse(invidivual.last_name)[0].inflect({'accs'}).word[1:]
    src = 'documents/dvictim.docx'
    dst = f'media/{cases.fk_user_id}/cases/{cases.active_cases_id}/documents/'
    shutil.copy(src, dst, follow_symlinks=True)
    if os.path.isfile(dst+'Постановление о признании потерпевшим.docx'):
        os.remove(dst+'Постановление о признании потерпевшим.docx')
    os.rename(dst+'dvictim.docx',dst+'Постановление о признании потерпевшим.docx')
    doc = docx.Document(dst+'Постановление о признании потерпевшим.docx')
    style = doc.styles['ConsNonformat']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    dictionary = {
        "App": user.appointment,
        "ur": user.users_rank.lower(),
        "FIO1": userFIO_1,
        "FIO2": userFIO_2,
        "(мс)": "г. "+department.city,
        "Домп": date_declaration,
        "rtle": declration.rationale,
        "toh": declration.type_of_harm.lower(),
        "vd": individual_dat,
        "vr": individual_rod,
        "numbcase": declarartion_criminal.number_cases
    }
    all_tables = doc.tables
    for i in dictionary:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i,dictionary[i])
        for k, table in enumerate(all_tables):
            for j, row in enumerate(table.rows):
                for cell in row.cells:
                    if cell.text.find(i) >= 0:
                        cell.text = cell.text.replace(i,dictionary[i])
    doc.save(dst+'Постановление о признании потерпевшим.docx')
    return FileResponse(dst+'Постановление о признании потерпевшим.docx',filename="Постановление о признании потерпевшим.docx",media_type='application/octet-stream')

@alg_router.get("/documents/declaration_criminal/{id}",tags=['Declaration Criminal Case'])
def get_declaration_criminal_case_document(id:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id==id).first()
    if cases is None:
        return JSONResponse(status_code=404,content={"message":"Дело не найдено"})
    statement = DB.query(Statement).filter(Statement.statement_id==cases.FK_statement_id).first()
    individual_id = statement.fk_individuals_id
    invidivual = DB.query(Invidivual).filter(Invidivual.individuals_id == individual_id).first()
    department_id = DB.query(User).filter(User.id==cases.fk_user_id).first().fk_department_id
    department = DB.query(Department).filter(Department.departments_id==department_id).first()
    declration = DB.query(DeclarationCriminalCase).filter(DeclarationCriminalCase.fk_active_case_id == cases.active_cases_id).first()
    if declration is None:
        return JSONResponse(status_code=404,content={"message":"Постановление не найдено"})
    date_declaration = declration.date_declaration.strftime("%d.%m.%Yг.")
    date_statement = statement.statement_date.strftime("%d.%m.%Y")
    date_procurator = declration.date_procurator.strftime("%d.%m.%Yг.")
    date_victim = declration.date_victim.strftime("%d.%m.%Yг.")
    morph = pymorphy3.MorphAnalyzer()
    date_attacker = ''
    fio_attacker = ''
    fullname_attacker = ''
    if declration.know_person == 2:
        date_attacker = ', а также '+declration.date_criminal_person.strftime("%d.%m.%Yг.")
        fullname_attacker = ', в отношении - '+ declration.firstname[0]+morph.parse(declration.firstname)[0].inflect({'accs'}).word[1:]+' '+declration.middlename[0]+morph.parse(declration.middlename)[0].inflect({'accs'}).word[1:]+' '+declration.lastname[0]+morph.parse(declration.lastname)[0].inflect({'accs'}).word[1:]
        fio_attacker = declration.firstname[0]+morph.parse(declration.firstname)[0].inflect({'datv'}).word[1:]+' '+declration.middlename[0]+'.'+declration.lastname[0] +'.'
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    userFIO_1 = user.first_name + ' ' + user.middle_name[0]+'.' + user.last_name[0]+'.'
    userFIO_2 = user.middle_name[0]+'.' + user.last_name[0]+'.'+ ' ' + user.first_name
    
    individual_rod = invidivual.first_name[0]+morph.parse(invidivual.first_name)[0].inflect({'accs'}).word[1:]+' '+ invidivual.middle_name[0]+morph.parse(invidivual.middle_name)[0].inflect({'accs'}).word[1:]+' '+ invidivual.last_name[0]+morph.parse(invidivual.last_name)[0].inflect({'accs'}).word[1:]
    fio_individual = invidivual.first_name[0]+morph.parse(invidivual.first_name)[0].inflect({'datv'}).word[1:]+' ' + invidivual.middle_name[0] +'.'+invidivual.last_name[0]+'.'
    src = 'documents/dcriminal.docx'
    dst = f'media/{cases.fk_user_id}/cases/{cases.active_cases_id}/documents/'
    shutil.copy(src, dst, follow_symlinks=True)
    if os.path.isfile(dst+'Постановление о возбуждении уголовного дела.docx'):
        os.remove(dst+'Постановление о возбуждении уголовного дела.docx')
    os.rename(dst+'dcriminal.docx',dst+'Постановление о возбуждении уголовного дела.docx')
    doc = docx.Document(dst+'Постановление о возбуждении уголовного дела.docx')
    style = doc.styles['ConsNonformat']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    dictionary = {
        "App": user.appointment,
        "ur": user.users_rank.lower(),
        "FIO1": userFIO_1,
        "FIO2": userFIO_2,
        "(мс)": "г. "+department.city,
        "Домп.": date_declaration,
        "infoc": statement.explanation_text,
        "dc": date_statement,
        "nd": department.name_department,
        "FIOZ" :individual_rod,
        "hd":declration.time_declaration[0] + declration.time_declaration[1],
        "md":declration.time_declaration[3] + declration.time_declaration[4],
        "argc":declration.reason_declaration,
        "its":declration.item_part_article,
        "np":declration.name_procurator,
        "dp":date_procurator,
        "ph":declration.time_procurator[0] + declration.time_procurator[1],
        "pm":declration.time_procurator[3] + declration.time_procurator[4],
        "dz":date_victim,
        "FIO3":fio_individual,
        "fa":fullname_attacker,
        "da":date_attacker,
        "FIO4":fio_attacker
    }
    all_tables = doc.tables
    for i in dictionary:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i,dictionary[i])
        for k, table in enumerate(all_tables):
            for j, row in enumerate(table.rows):
                for cell in row.cells:
                    if cell.text.find(i) >= 0:
                        cell.text = cell.text.replace(i,dictionary[i])
    doc.save(dst+'Постановление о возбуждении уголовного дела.docx')
    return FileResponse(dst+'Постановление о возбуждении уголовного дела.docx',filename="Постановление о возбуждении уголовного дела.docx",media_type='application/octet-stream')
@alg_router.get("/documents/order_wanted/{id}",tags=['Order Wanted'])
def get_declaration_order_wanted(id:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id==id).first()
    if cases is None:
        return JSONResponse(status_code=404,content={"message":"Дело не найдено"})
    statement = DB.query(Statement).filter(Statement.statement_id==cases.FK_statement_id).first()
    individual_id = statement.fk_individuals_id
    invidivual = DB.query(Invidivual).filter(Invidivual.individuals_id == individual_id).first()
    department_id = DB.query(User).filter(User.id==cases.fk_user_id).first().fk_department_id
    department = DB.query(Department).filter(Department.departments_id==department_id).first()
    order_wanted = DB.query(OrderWanted).filter(OrderWanted.fk_active_case_id == cases.active_cases_id).first()
    if order_wanted is None:
        return JSONResponse(status_code=404,content={"message":"Поручение не найдено"})
    date_statement = statement.statement_date.strftime("%d.%m.%Y")
    morph = pymorphy3.MorphAnalyzer()
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    userFIO_2 = user.middle_name[0]+'.' + user.last_name[0]+'.'+ ' ' + user.first_name
    individual_rod = invidivual.first_name[0]+morph.parse(invidivual.first_name)[0].inflect({'accs'}).word[1:]+' '+ invidivual.middle_name[0]+morph.parse(invidivual.middle_name)[0].inflect({'accs'}).word[1:]+' '+ invidivual.last_name[0]+morph.parse(invidivual.last_name)[0].inflect({'accs'}).word[1:]
    src = 'documents/ow.docx'
    dst = f'media/{cases.fk_user_id}/cases/{cases.active_cases_id}/documents/'
    shutil.copy(src, dst, follow_symlinks=True)
    if os.path.isfile(dst+'Поручение о проведении оперативно-розыскных мероприятий.docx'):
        os.remove(dst+'Поручение о проведении оперативно-розыскных мероприятий.docx')
    os.rename(dst+'ow.docx',dst+'Поручение о проведении оперативно-розыскных мероприятий.docx')
    doc = docx.Document(dst+'Поручение о проведении оперативно-розыскных мероприятий.docx')
    style = doc.styles['ConsNonformat']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    dictionary = {
        "App": user.appointment,
        "ur": user.users_rank.lower(),
        "FIO2": userFIO_2,
        "dc": date_statement,
        "FIOZ" :individual_rod,
        "nd": department.name_department,
        "infoc":statement.explanation_text.lower(),
        "chc":order_wanted.check_criminal,
        "ow":order_wanted.text_order_wanted,
        "N_Or":order_wanted.name_orders
    }
    all_tables = doc.tables
    for i in dictionary:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i,dictionary[i])
        for k, table in enumerate(all_tables):
            for j, row in enumerate(table.rows):
                for cell in row.cells:
                    if cell.text.find(i) >= 0:
                        cell.text = cell.text.replace(i,dictionary[i])
    doc.save(dst+'Поручение о проведении оперативно-розыскных мероприятий.docx')
    return FileResponse(dst+'Поручение о проведении оперативно-розыскных мероприятий.docx',filename="Поручение о проведении оперативно-розыскных мероприятий.docx",media_type='application/octet-stream')

@alg_router.get("/documents/clarif/{id}",tags=['Clarif'])
def get_clarif_document(id:int,DB:Session = Depends(get_session)):
    questperson = DB.query(QuestPerson).filter(QuestPerson.id_quest_persons == id).first()
    if questperson is None:
        return JSONResponse(status_code=404,content={"message":"Допрашиваемый человек не найден"})
    questperson_individual = DB.query(Invidivual).filter(Invidivual.individuals_id == questperson.individuals_id).first()
    passport = DB.query(Passport).filter(Passport.passports_id == questperson_individual.FK_passports_id).first()
    cases = DB.query(Case).filter(Case.active_cases_id==questperson.fk_active_case_id).first()
    if cases is None:
        return JSONResponse(status_code=404,content={"message":"Дело не найдено"})
    department_id = DB.query(User).filter(User.id==cases.fk_user_id).first().fk_department_id
    department = DB.query(Department).filter(Department.departments_id==department_id).first()
    clarif = DB.query(Clarif).filter(Clarif.fk_id_quest_person == questperson.id_quest_persons).first()
    if clarif is None:
        return JSONResponse(status_code=404,content={"message":"Объяснение не найдено"})
    true_clarif = "правильно"
    if clarif.true_clarif == 2:
        true_clarif = "не правильно"
    addition_clarif = 'отсутсвуют'
    if len(clarif.addition_clarif) !=0:
        addition_clarif = clarif.addition_clarif
    date_clarif = clarif.date_clarif.strftime("%d.%m.%Y")
    birthday = passport.birthdate.strftime("%d.%m.%Y")
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    userFIO_1 = user.first_name + ' ' + user.middle_name[0]+'.' + user.last_name[0]+'.'
    userFIO_2 = user.middle_name[0]+'.' + user.last_name[0]+'.'+ ' ' + user.first_name
    quest_people = DB.query(ClarifPersonInvolved).filter(clarif.id_clarif == ClarifPersonInvolved.id_clarif).all()
    morph = pymorphy3.MorphAnalyzer()
    otherPerson = ''
    for i in range(len(quest_people)):
        person = DB.query(Invidivual).filter(Invidivual.individuals_id == quest_people[i].id_persons_involved).first()
        procedural_position_name = DB.query(ProceduralPosition).filter(quest_people[i].fk_procedural_position==ProceduralPosition.id_procedural_position).first().name_position
        if procedural_position_name != 'Потерпевший' and procedural_position_name !='Свидетель' and procedural_position_name !='Подозреваемый':
            if len(otherPerson) == 0:
                otherPerson = ', с участием '
            elif len(otherPerson) > 0:
                otherPerson += ', '
            otherPerson += morph.parse(procedural_position_name.lower())[0].inflect({'gent'}).word + ' '+ person.first_name[0]+morph.parse(person.first_name)[0].inflect({'gent'}).word[1:] + ' ' + person.middle_name[0]+morph.parse(person.middle_name)[0].inflect({'gent'}).word[1:] + ' ' + person.last_name[0]+morph.parse(person.last_name)[0].inflect({'gent'}).word[1:]
    dst = f'media/{cases.fk_user_id}/cases/{cases.active_cases_id}/documents/'    
    src = 'documents/clarif.docx'
    dstsrc = 'Объяснение от потерпевшего.docx'
    shutil.copy(src, dst, follow_symlinks=True)
    if os.path.isfile(dst+'Объяснение от потерпевшего.docx'):
        os.remove(dst+'Объяснение от потерпевшего.docx')
    os.rename(dst+'clarif.docx',dst+'Объяснение от потерпевшего.docx')
    doc = docx.Document(dst+dstsrc)
    style1 = doc.styles['Normal']
    style1.font.name = 'Times New Roman'
    style1.font.size = Pt(13)
    style = doc.styles['ConsNonformat']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    dictionary = {
        "fiovictim": questperson_individual.first_name +' ' + questperson_individual.middle_name+' ' + questperson_individual.last_name,
        "birthvictim": birthday,
        "placebirth": passport.place_of_birth,
        "resvictim": passport.place_of_residence,
        "publvic": questperson_individual.resident,
        "educvic": questperson_individual.education,
        "falmvic": questperson_individual.family_status,
        "workvic": questperson_individual.place_work,
        "warvic": questperson_individual.military,
        "jailvic": questperson_individual.criminal,
        "passpvic": passport.passport_serial + ' '+ passport.passport_number,
        "othervic": questperson_individual.other,
        "othp": otherPerson,
        "App": user.appointment,
        "ur": user.users_rank.lower(),
        "FIO1": userFIO_1,
        "FIO2": userFIO_2,
        "(мс)": "г. "+department.city,
        "Домп.": date_clarif,
        "whc": clarif.place_clarif,
        "bch": clarif.begin_clarif[0]+clarif.begin_clarif[1],
        "bcm":clarif.begin_clarif[3]+clarif.begin_clarif[4],
        "ech": clarif.end_clarif[0]+clarif.end_clarif[1],
        "ecm":clarif.end_clarif[3]+clarif.end_clarif[4],
        "tm": clarif.technical_means.lower(),
        "clarif": clarif.text_clarif,
        "cr": clarif.read_clarif,
        "tf":true_clarif,
        "dop": addition_clarif,
        "ph":clarif.photography.lower()
    }
    all_tables = doc.tables
    
    for i in dictionary:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i,dictionary[i])
        for k, table in enumerate(all_tables):
            for j, row in enumerate(table.rows):
                for cell in row.cells:
                    if cell.text.find(i) >= 0:
                        cell.text = cell.text.replace(i,dictionary[i])
    doc.save(dst+dstsrc)
    return FileResponse(dst+dstsrc,filename="Объяснение от потерпевшего.docx",media_type='application/octet-stream')

@alg_router.get("/documents/get_all_requests/{id}", tags=['Request'])
def get_all_requests_docs_names(id:int, DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    if os.path.exists(f'media/{user.id}/cases/{cases.active_cases_id}') == False:
        return JSONResponse(status_code= 404, content={"message": "Добавление файлов невозможно!"})
    files = os.listdir(f'media/{user.id}/cases/{cases.active_cases_id}/documents')
    if len(files) == 0:
        return JSONResponse(status_code=404,content={"message":"Файлы не найдены"})
    request_files = []
    for file in files:
        if file[:6] == 'Запрос':
            request_files.append(file)
    if len(request_files) == 0:
        return JSONResponse(status_code=404,content={"message":"Файлы не найдены"})
    return request_files
@alg_router.get("/documents/get_all_docs/{id}",tags=['Documents'])
def get_all_documents_names(id:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    if os.path.exists(f'media/{user.id}/cases/{cases.active_cases_id}') == False:
        return JSONResponse(status_code= 404, content={"message": "Добавление файлов невозможно!"})
    files = os.listdir(f'media/{user.id}/cases/{cases.active_cases_id}/documents')
    if len(files) == 0:
        return JSONResponse(status_code=404,content={"message":"Файлы не найдены"})
    docs_files = []
    for file in files:
        if file[:6] != 'Запрос' and file[:6] != 'Допрос':
            docs_files.append(file)
    if len(docs_files) == 0:
        return JSONResponse(status_code=404,content={"message":"Файлы не найдены"})
    return docs_files
@alg_router.get("/documents/get_all_quest/{id}",tags=['Quest'])
def get_all_quest_names(id:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    if os.path.exists(f'media/{user.id}/cases/{cases.active_cases_id}') == False:
        return JSONResponse(status_code= 404, content={"message": "Добавление файлов невозможно!"})
    files = os.listdir(f'media/{user.id}/cases/{cases.active_cases_id}/documents')
    if len(files) == 0:
        return JSONResponse(status_code=404,content={"message":"Файлы не найдены"})
    quest_files = []
    for file in files:
        if file[:6] == 'Допрос':
            quest_files.append(file)
    if len(quest_files) == 0:
        return JSONResponse(status_code=404,content={"message":"Файлы не найдены"})
    return quest_files
@alg_router.get("/documents/{id}/download/{filenames}",tags=['Documents'])
def download_passport_file(id:int,filenames:str,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    if os.path.exists(f'media/{user.id}/cases/{cases.active_cases_id}/documents/{filenames}') == False:
        return JSONResponse(status_code= 404, content={"message": "Файл не существует"})
    return FileResponse(f'media/{user.id}/cases/{cases.active_cases_id}/documents/{filenames}',filename=f'{filenames}',media_type='application/octet-stream')